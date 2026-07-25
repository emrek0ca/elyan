from __future__ import annotations

import os
import tempfile
import uuid
from pathlib import Path
from typing import Any

from actions._read_only_common import bulletize_text, ensure_allowed_path, summarize_text
from actions._visual_block_common import (
    normalize_visual_blocks,
    render_chart_block_to_png,
    split_text_blocks,
    table_ensure_width,
    validate_visual_block_paths,
)
from actions._write_common import artifact_payload, ensure_allowed_output_path, normalize_source_context, sanitize_xml_text
from actions.document_read import _ALLOWED_SUFFIXES as _READ_ALLOWED_SUFFIXES
from actions.document_read import _extract_document_text
from runtime.capability_registry import SafeCapabilityError


def _workspace_root() -> Path:
    from actions._read_only_common import workspace_root

    return workspace_root()


def _chart_temp_path() -> Path:
    directory = Path(tempfile.gettempdir())
    directory.mkdir(parents=True, exist_ok=True)
    fd, raw_path = tempfile.mkstemp(prefix=f"elyan-doc-chart-{uuid.uuid4().hex[:8]}-", suffix=".png", dir=str(directory))
    os.close(fd)
    return Path(raw_path)


def _resolve_source_text(
    source_path: str,
    source_context: str,
    prompt: str,
    *,
    allow_empty: bool = False,
    selected_paths: list[str] | None = None,
) -> tuple[str, str]:
    if str(source_context or "").strip():
        text = sanitize_xml_text(str(source_context or "").strip(), max_chars=60000)
        return text, normalize_source_context(text)
    if str(source_path or "").strip():
        resolved = ensure_allowed_path(
            source_path,
            allowed_suffixes=_READ_ALLOWED_SUFFIXES,
            selected_paths=selected_paths,
            root_resolver=_workspace_root,
        )
        extracted, _ = _extract_document_text(resolved)
        text = sanitize_xml_text(str(extracted or "").strip(), max_chars=60000)
        if not text:
            raise SafeCapabilityError("EMPTY_DOCUMENT", "Kaynak belgede yazılabilir içerik bulunamadı.")
        return text, resolved.name
    text = sanitize_xml_text(str(prompt or "").strip(), max_chars=60000)
    if not text:
        if allow_empty:
            return "", normalize_source_context(source_context or prompt or "structured document")
        raise SafeCapabilityError("INVALID_ARGUMENT", "Belge oluşturmak için bir içerik veya hedef gerekli.")
    return text, normalize_source_context(text)


def _section_blocks(sections: list[dict[str, Any]] | None) -> list[dict[str, Any]]:
    blocks: list[dict[str, Any]] = []
    for section in sections or []:
        if not isinstance(section, dict):
            continue
        heading = str(section.get("heading", "") or section.get("title", "") or "").strip()
        body = str(section.get("body", "") or section.get("text", "") or "").strip()
        if heading:
            blocks.append({"kind": "text", "text": heading, "level": 1})
        if body:
            blocks.extend(split_text_blocks(body))
        nested_blocks = section.get("blocks", [])
        if isinstance(nested_blocks, list) and nested_blocks:
            blocks.extend(normalize_visual_blocks(nested_blocks, fallback_text=""))
    return blocks


def _add_table(document: Any, block: dict[str, Any]) -> None:
    headers, rows = table_ensure_width(
        [str(item) for item in (block.get("headers", []) or [])],
        [[str(cell) for cell in row] for row in (block.get("rows", []) or [])],
    )
    if not rows and not headers:
        return
    if str(block.get("title", "") or "").strip():
        document.add_paragraph(str(block.get("title")).strip())
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, header in enumerate(headers):
        header_cells[index].text = header
    for row in rows:
        new_row = table.add_row().cells
        for index, cell in enumerate(row):
            new_row[index].text = cell


def _add_image(document: Any, block: dict[str, Any]) -> None:
    from docx.shared import Inches  # type: ignore[reportMissingImports]

    path = Path(str(block.get("path", "") or "")).expanduser()
    if not path.exists():
        return
    picture_width = float(block.get("width", 0) or 0)
    picture_height = float(block.get("height", 0) or 0)
    if picture_width > 0:
        width = Inches(max(1.5, min(6.5, picture_width / 144.0)))
        document.add_picture(str(path), width=width)
    elif picture_height > 0:
        height = Inches(max(1.0, min(6.5, picture_height / 144.0)))
        document.add_picture(str(path), height=height)
    else:
        document.add_picture(str(path), width=Inches(5.5))


def _add_blocks(document: Any, blocks: list[dict[str, Any]], temp_paths: list[Path] | None = None) -> None:
    for block in blocks:
        kind = str(block.get("kind", "") or "").strip().lower()
        if kind == "text":
            text = sanitize_xml_text(str(block.get("text", "") or "").strip(), max_chars=60000)
            if not text:
                continue
            level = int(block.get("level", 0) or 0)
            if level > 0:
                document.add_heading(text, level=min(level, 3))
            else:
                for paragraph in [part.strip() for part in text.split("\n\n") if part.strip()]:
                    document.add_paragraph(paragraph)
            continue
        if kind == "table":
            _add_table(document, block)
            continue
        if kind == "image":
            _add_image(document, block)
            continue
        if kind == "chart":
            chart_path = _chart_temp_path()
            try:
                render_chart_block_to_png(block, chart_path)
                if temp_paths is not None:
                    temp_paths.append(chart_path)
                _add_image(document, {"path": str(chart_path), "width": block.get("width", 0), "height": block.get("height", 0)})
            except Exception:
                try:
                    chart_path.unlink(missing_ok=True)
                except Exception:
                    pass
            continue
        if kind == "spacer":
            document.add_paragraph("")


def _collect_citations(kwargs: dict[str, Any]) -> list[dict[str, str]]:
    """Atıfları argümanlardan ve bağımlılık sonuçlarından toplar.

    Kaynaklar iki yoldan gelir: doğrudan ``citations`` argümanı ya da önceki
    adımların (web araştırması / Compound) yapılandırılmış çıktısı. URL'e göre
    tekilleştirilir."""
    found: list[dict[str, str]] = []

    def _absorb(value: Any) -> None:
        if not isinstance(value, list):
            return
        for item in value:
            if not isinstance(item, dict):
                continue
            url = str(item.get("url") or item.get("link") or "").strip()
            if not url:
                continue
            title = str(item.get("title") or item.get("name") or "").strip() or url
            found.append({"title": title[:200], "url": url[:500]})

    _absorb(kwargs.get("citations"))
    _absorb(kwargs.get("sources"))
    for payload in (kwargs.get("_dependencyResults") or {}).values() if isinstance(
        kwargs.get("_dependencyResults"), dict
    ) else ():
        if not isinstance(payload, dict):
            continue
        result = payload.get("result") if isinstance(payload.get("result"), dict) else payload
        for key in ("citations", "sources", "results", "webSources"):
            _absorb(result.get(key) if isinstance(result, dict) else None)

    seen: set[str] = set()
    unique: list[dict[str, str]] = []
    for item in found:
        if item["url"] in seen:
            continue
        seen.add(item["url"])
        unique.append(item)
    return unique[:20]


def _add_sources_section(document: Any, citations: list[dict[str, str]]) -> None:
    """Belgenin sonuna numaralı 'Kaynaklar' bölümü ekler (atıf varsa)."""
    if not citations:
        return
    document.add_paragraph("")
    document.add_heading("Kaynaklar", level=1)
    for index, citation in enumerate(citations, start=1):
        document.add_paragraph(
            f"{index}. {sanitize_xml_text(citation['title'], max_chars=200)} — {citation['url']}"
        )


def document_write(
    prompt: str = "",
    output_path: str = "",
    title: str = "",
    sections: list[dict[str, Any]] | None = None,
    blocks: list[dict[str, Any]] | None = None,
    source_path: str = "",
    source_context: str = "",
    overwrite: bool = False,
    _selectedPaths: list[str] | None = None,
    **kwargs: Any,
) -> dict[str, Any]:
    from docx import Document  # type: ignore[reportMissingImports]

    if not str(prompt or "").strip():
        for key in ("content", "body", "text", "markdown", "instructions", "instruction", "description"):
            candidate = kwargs.get(key)
            if str(candidate or "").strip():
                prompt = str(candidate)
                break
    document_payload = kwargs.get("document") if isinstance(kwargs.get("document"), dict) else kwargs.get("doc")
    if isinstance(document_payload, dict):
        if not str(prompt or "").strip():
            for key in ("prompt", "content", "body", "text", "markdown", "summary"):
                candidate = document_payload.get(key)
                if str(candidate or "").strip():
                    prompt = str(candidate)
                    break
        if not sections and isinstance(document_payload.get("sections"), list):
            sections = document_payload.get("sections")
        if not blocks and isinstance(document_payload.get("blocks"), list):
            blocks = document_payload.get("blocks")
    if not str(output_path or "").strip():
        output_path = str(kwargs.get("outputPath", "") or kwargs.get("output_path", "") or "")
    if not str(source_context or "").strip():
        source_context = str(kwargs.get("sourceContext", "") or kwargs.get("source_context", "") or "")

    resolved_output = ensure_allowed_output_path(
        output_path,
        extension=".docx",
        overwrite=overwrite,
        hint=title or prompt or source_path or "elyan-document",
        root_resolver=_workspace_root,
    )
    body_text, source_summary = _resolve_source_text(
        source_path,
        source_context,
        prompt,
        allow_empty=bool((blocks or []) or (sections or [])),
        selected_paths=_selectedPaths,
    )

    document = Document()
    heading = str(title or resolved_output.stem).strip() or "Elyan Document"
    document.add_heading(heading, level=0)
    temp_paths: list[Path] = []

    try:
        normalized_blocks = normalize_visual_blocks(
            [*(blocks or []), *_section_blocks([item for item in (sections or []) if isinstance(item, dict)])],
            fallback_text=body_text,
        )
        normalized_blocks = validate_visual_block_paths(
            normalized_blocks,
            selected_paths=_selectedPaths,
            root_resolver=_workspace_root,
        )
        _add_blocks(document, normalized_blocks, temp_paths=temp_paths)
        # Kaynaklar belgenin İÇİNE yazılır. Atıf bir sohbet süsü değil, çıktının
        # denetlenebilirliğidir: kullanıcı belgeyi paylaştığında dayanağı da
        # yanında gider. Atıf yoksa bölüm hiç eklenmez (boş başlık bırakılmaz).
        _add_sources_section(document, _collect_citations(kwargs))

        document.save(str(resolved_output))
    finally:
        for path in temp_paths:
            try:
                path.unlink(missing_ok=True)
            except Exception:
                continue
    summary = summarize_text(body_text, max_chars=260)
    return {
        "text": f"DOCX oluşturuldu: {resolved_output.name}",
        "result": {
            "kind": "document_write",
            "sourceContext": source_summary,
            "sourcePath": str(source_path or "").strip(),
            "outputPath": str(resolved_output),
            "contentType": artifact_payload(resolved_output)["contentType"],
            "created": True,
            "title": heading,
            "summary": summary,
            "bullets": bulletize_text(body_text),
        },
        "artifacts": [artifact_payload(resolved_output)],
    }
