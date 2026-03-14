"""Word Document Tools - Read and Write .docx files"""

import asyncio
from pathlib import Path
from typing import Any
from utils.logger import get_logger
from security.validator import validate_path
from config.settings import HOME_DIR, DESKTOP

logger = get_logger("office.word")

# Maximum file size (10MB)
MAX_FILE_SIZE = 10 * 1024 * 1024


async def read_word(path: str, max_chars: int = 10000) -> dict[str, Any]:
    """Read content from a Word document (.docx)

    Args:
        path: Path to the Word document
        max_chars: Maximum characters to return (default 10000)
    """
    try:
        # Validate path
        is_valid, error, _ = validate_path(path)
        if not is_valid:
            return {"success": False, "error": error}

        file_path = Path(path).expanduser().resolve()

        if not file_path.exists():
            return {"success": False, "error": f"Dosya bulunamadı: {file_path.name}"}

        if not file_path.suffix.lower() in [".docx", ".doc"]:
            return {"success": False, "error": "Sadece .docx dosyaları destekleniyor"}

        # Check file size
        if file_path.stat().st_size > MAX_FILE_SIZE:
            return {"success": False, "error": "Dosya çok büyük (max 10MB)"}

        # Import here to avoid startup delay if not used
        try:
            from docx import Document
        except ImportError:
            return {"success": False, "error": "python-docx kurulu değil. 'pip install python-docx' çalıştırın."}

        # Run in thread pool to avoid blocking
        def _read():
            doc = Document(str(file_path))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also get text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        tables_text.append(row_text)

            return paragraphs, tables_text, len(doc.paragraphs), len(doc.tables)

        loop = asyncio.get_event_loop()
        paragraphs, tables_text, para_count, table_count = await loop.run_in_executor(None, _read)

        # Combine content
        content = "\n\n".join(paragraphs)
        if tables_text:
            content += "\n\n--- Tablolar ---\n" + "\n".join(tables_text)

        # Truncate if needed
        truncated = False
        if len(content) > max_chars:
            content = content[:max_chars]
            truncated = True

        logger.info(f"Read Word file: {file_path.name}, {len(content)} chars")

        return {
            "success": True,
            "path": str(file_path),
            "filename": file_path.name,
            "content": content,
            "paragraph_count": para_count,
            "table_count": table_count,
            "truncated": truncated
        }

    except Exception as e:
        logger.error(f"Read Word error: {e}")
        return {"success": False, "error": str(e)}


async def write_word(
    path: str = None,
    content: str = "",
    title: str = None,
    paragraphs: list = None
) -> dict[str, Any]:
    """Create or write to a Word document (.docx)

    Args:
        path: Path for the document (default: Desktop/document.docx)
        content: Text content to write (will be split by newlines into paragraphs)
        title: Optional document title (added as heading)
        paragraphs: Optional list of paragraphs (alternative to content string)
    """
    try:
        # Default path
        if not path:
            path = str(DESKTOP / "belge.docx")

        # Ensure .docx extension
        file_path = Path(path).expanduser().resolve()
        if not file_path.suffix.lower() == ".docx":
            file_path = file_path.with_suffix(".docx")

        # Validate path
        is_valid, error, _ = validate_path(str(file_path))
        if not is_valid:
            return {"success": False, "error": error}

        content_text = str(content or "").strip()
        paragraph_items: list[str] = []
        if isinstance(paragraphs, list):
            for para in paragraphs:
                para_text = str(para or "").strip()
                if para_text:
                    paragraph_items.append(para_text)

        # Rule-1: No empty or skeleton files
        total_len = len(content_text) + sum(len(p) for p in paragraph_items) + len(str(title or ""))
        if total_len == 0:
            return {
                "success": False,
                "error": "EMPTY_CONTENT: Word içeriği boş. Lütfen yazılacak bir içerik verin.",
                "error_code": "EMPTY_CONTENT",
            }
        if total_len < 200:
            return {
                "success": False, 
                "error": f"CONTENT_TOO_SHORT: Word içeriği çok kısa ({total_len} karakter). En az 200 karakter olmalı.",
                "error_code": "CONTENT_TOO_SHORT"
            }

        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches, Pt
        except ImportError:
            return {"success": False, "error": "python-docx kurulu değil. 'pip install python-docx' çalıştırın."}

        def _write():
            doc = Document()
            try:
                for section in doc.sections:
                    section.top_margin = Inches(0.8)
                    section.bottom_margin = Inches(0.8)
                    section.left_margin = Inches(0.85)
                    section.right_margin = Inches(0.85)
            except Exception:
                pass

            try:
                normal_style = doc.styles["Normal"]
                normal_style.font.name = "Aptos"
                normal_style.font.size = Pt(11)
            except Exception:
                pass

            heading_markers = {
                "Kısa Özet",
                "Temel Bulgular",
                "Kaynak Güven Özeti",
                "Açık Riskler",
                "Belirsizlikler",
                "Kaynakça",
                "İçerik",
            }

            # Add title if provided
            if title:
                title_p = doc.add_paragraph()
                title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = title_p.add_run(str(title).strip())
                run.bold = True
                run.font.size = Pt(18)
                title_p.paragraph_format.space_after = Pt(18)

            # Add content
            if paragraph_items:
                for para in paragraph_items:
                    para_text = str(para or "").strip()
                    if not para_text:
                        continue
                    if para_text in heading_markers:
                        heading = doc.add_paragraph()
                        heading_run = heading.add_run(para_text)
                        heading_run.bold = True
                        heading_run.font.size = Pt(13)
                        heading.paragraph_format.space_before = Pt(10)
                        heading.paragraph_format.space_after = Pt(6)
                    else:
                        body = doc.add_paragraph()
                        body.add_run(para_text)
                        body.paragraph_format.space_after = Pt(8)
            elif content_text:
                # Split by double newlines for paragraphs
                for para in content_text.split("\n\n"):
                    if para.strip():
                        body = doc.add_paragraph()
                        body.add_run(para.strip())
                        body.paragraph_format.space_after = Pt(8)

            # Create parent directory if needed
            file_path.parent.mkdir(parents=True, exist_ok=True)

            doc.save(str(file_path))
            return True

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, _write)

        # Post-check validation
        if not file_path.exists():
            return {"success": False, "error": "WRITE_FAILED: Dosya diskte bulunamadı.", "error_code": "FILE_NOT_FOUND"}
        
        file_size = file_path.stat().st_size
        if file_size < 2000: # Minimum docx structure size + small content
            return {
                "success": False, 
                "error": f"WRITE_POSTCHECK_FAILED: Dosya boyutu çok küçük ({file_size} bytes). İçerik tam yazılamamış olabilir.",
                "error_code": "WRITE_POSTCHECK_FAILED"
            }

        logger.info(f"Created Word file: {file_path.name} ({file_size} bytes)")

        return {
            "success": True,
            "path": str(file_path),
            "filename": file_path.name,
            "size_bytes": file_size
        }

    except Exception as e:
        logger.error(f"Write Word error: {e}")
        return {"success": False, "error": str(e)}
