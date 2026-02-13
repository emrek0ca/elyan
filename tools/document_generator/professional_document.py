"""
Professional Document Generator
Creates high-quality documents from research results
Supports multiple formats: PDF, DOCX, HTML, Markdown
"""

import json
import os
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional
from dataclasses import dataclass
from enum import Enum
from utils.logger import get_logger

logger = get_logger("document_generator")


class DocumentFormat(Enum):
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    MARKDOWN = "markdown"
    TXT = "txt"


class DocumentTemplate(Enum):
    RESEARCH_REPORT = "research_report"
    EXECUTIVE_SUMMARY = "executive_summary"
    ACADEMIC_PAPER = "academic_paper"
    BUSINESS_REPORT = "business_report"
    TECHNICAL_DOCUMENTATION = "technical_documentation"
    PRESENTATION_NOTES = "presentation_notes"


@dataclass
class DocumentSection:
    """Represents a document section"""
    title: str
    content: str
    level: int = 1  # Heading level
    subsections: List['DocumentSection'] = None

    def __post_init__(self):
        if self.subsections is None:
            self.subsections = []


@dataclass
class DocumentMetadata:
    """Document metadata"""
    title: str
    author: str = "Elyan AI Assistant"
    date: str = None
    version: str = "1.0"
    language: str = "tr"
    keywords: List[str] = None
    abstract: str = ""

    def __post_init__(self):
        if self.date is None:
            self.date = datetime.now().strftime("%Y-%m-%d")
        if self.keywords is None:
            self.keywords = []


class ProfessionalDocumentGenerator:
    """Generate professional documents from research and data"""

    def __init__(self, output_dir: str = None):
        self.output_dir = Path(output_dir) if output_dir else Path.home() / "Desktop" / "ElyanDocuments"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def generate_research_report(
        self,
        research_data: Dict[str, Any],
        template: DocumentTemplate = DocumentTemplate.RESEARCH_REPORT,
        format: DocumentFormat = DocumentFormat.DOCX,
        custom_title: str = None,
        include_toc: bool = True,
        include_bibliography: bool = True,
        language: str = "tr"
    ) -> Dict[str, Any]:
        """
        Generate a research report from research data

        Args:
            research_data: Research results from deep_research
            template: Document template to use
            format: Output format
            custom_title: Custom document title
            include_toc: Include table of contents
            include_bibliography: Include bibliography
            language: Document language

        Returns:
            Result with file path and metadata
        """
        try:
            # Extract data
            topic = research_data.get("topic", "Research Report")
            summary = research_data.get("summary", "")
            findings = research_data.get("findings", [])
            key_insights = research_data.get("key_insights", [])
            sources = research_data.get("sources", [])
            statistics = research_data.get("statistics", {})
            bibliography = research_data.get("bibliography", [])

            # Create metadata
            metadata = DocumentMetadata(
                title=custom_title or f"{topic} - Araştırma Raporu",
                language=language,
                keywords=self._extract_keywords(findings),
                abstract=summary[:500] if summary else ""
            )

            # Build sections
            sections = self._build_report_sections(
                topic=topic,
                summary=summary,
                findings=findings,
                key_insights=key_insights,
                sources=sources,
                statistics=statistics,
                bibliography=bibliography if include_bibliography else [],
                language=language
            )

            # Generate document
            if format == DocumentFormat.DOCX:
                return self._generate_docx(metadata, sections, include_toc)
            elif format == DocumentFormat.PDF:
                return self._generate_pdf(metadata, sections, include_toc)
            elif format == DocumentFormat.HTML:
                return self._generate_html(metadata, sections, include_toc)
            elif format == DocumentFormat.MARKDOWN:
                return self._generate_markdown(metadata, sections)
            else:
                return self._generate_txt(metadata, sections)

        except Exception as e:
            logger.error(f"Document generation error: {e}")
            return {
                "success": False,
                "error": str(e)
            }

    def _build_report_sections(
        self,
        topic: str,
        summary: str,
        findings: List[Dict],
        key_insights: List[str],
        sources: List[Dict],
        statistics: Dict,
        bibliography: List[str],
        language: str = "tr"
    ) -> List[DocumentSection]:
        """Build report sections"""
        sections = []

        # Labels based on language
        if language == "tr":
            labels = {
                "executive_summary": "Yönetici Özeti",
                "introduction": "Giriş",
                "key_findings": "Temel Bulgular",
                "detailed_analysis": "Detaylı Analiz",
                "statistics": "İstatistikler",
                "sources": "Kaynaklar",
                "bibliography": "Kaynakça",
                "conclusion": "Sonuç",
                "methodology": "Metodoloji"
            }
        else:
            labels = {
                "executive_summary": "Executive Summary",
                "introduction": "Introduction",
                "key_findings": "Key Findings",
                "detailed_analysis": "Detailed Analysis",
                "statistics": "Statistics",
                "sources": "Sources",
                "bibliography": "Bibliography",
                "conclusion": "Conclusion",
                "methodology": "Methodology"
            }

        # Executive Summary
        if summary:
            sections.append(DocumentSection(
                title=labels["executive_summary"],
                content=summary,
                level=1
            ))

        # Key Insights
        if key_insights:
            insights_content = "\n".join([f"• {insight}" for insight in key_insights[:10]])
            sections.append(DocumentSection(
                title=labels["key_findings"],
                content=insights_content,
                level=1
            ))

        # Detailed Findings by Category
        if findings:
            categories = {}
            for finding in findings:
                cat = finding.get("category", "general")
                if cat not in categories:
                    categories[cat] = []
                categories[cat].append(finding)

            category_labels = {
                "definition": "Tanımlar" if language == "tr" else "Definitions",
                "statistics": "İstatistikler" if language == "tr" else "Statistics",
                "research": "Araştırma Bulguları" if language == "tr" else "Research Findings",
                "expert_opinion": "Uzman Görüşleri" if language == "tr" else "Expert Opinions",
                "historical": "Tarihsel Bilgi" if language == "tr" else "Historical Information",
                "general": "Genel Bilgiler" if language == "tr" else "General Information"
            }

            detailed_section = DocumentSection(
                title=labels["detailed_analysis"],
                content="",
                level=1
            )

            for cat, cat_findings in categories.items():
                cat_content = "\n\n".join([
                    f"• {f.get('content', '')}" for f in cat_findings[:5]
                ])
                detailed_section.subsections.append(DocumentSection(
                    title=category_labels.get(cat, cat.title()),
                    content=cat_content,
                    level=2
                ))

            sections.append(detailed_section)

        # Statistics Section
        if statistics:
            stats_content = self._format_statistics(statistics, language)
            sections.append(DocumentSection(
                title=labels["statistics"],
                content=stats_content,
                level=1
            ))

        # Sources Overview
        if sources:
            sources_content = self._format_sources(sources, language)
            sections.append(DocumentSection(
                title=labels["sources"],
                content=sources_content,
                level=1
            ))

        # Bibliography
        if bibliography:
            bib_content = "\n".join([f"{i+1}. {bib}" for i, bib in enumerate(bibliography)])
            sections.append(DocumentSection(
                title=labels["bibliography"],
                content=bib_content,
                level=1
            ))

        return sections

    def _format_statistics(self, statistics: Dict, language: str) -> str:
        """Format statistics for display"""
        lines = []

        if language == "tr":
            if "total_sources" in statistics:
                lines.append(f"• Toplam Kaynak Sayısı: {statistics['total_sources']}")
            if "total_findings" in statistics:
                lines.append(f"• Toplam Bulgu Sayısı: {statistics['total_findings']}")
            if "avg_reliability" in statistics:
                lines.append(f"• Ortalama Güvenilirlik: %{statistics['avg_reliability']*100:.1f}")
            if "high_importance_findings" in statistics:
                lines.append(f"• Yüksek Önemli Bulgular: {statistics['high_importance_findings']}")

            if "source_types" in statistics:
                lines.append("\nKaynak Türleri:")
                st = statistics["source_types"]
                if st.get("academic", 0) > 0:
                    lines.append(f"  - Akademik: {st['academic']}")
                if st.get("news", 0) > 0:
                    lines.append(f"  - Haber: {st['news']}")
                if st.get("wiki", 0) > 0:
                    lines.append(f"  - Vikipedi: {st['wiki']}")
                if st.get("web", 0) > 0:
                    lines.append(f"  - Web: {st['web']}")
        else:
            if "total_sources" in statistics:
                lines.append(f"• Total Sources: {statistics['total_sources']}")
            if "total_findings" in statistics:
                lines.append(f"• Total Findings: {statistics['total_findings']}")
            if "avg_reliability" in statistics:
                lines.append(f"• Average Reliability: {statistics['avg_reliability']*100:.1f}%")

        return "\n".join(lines)

    def _format_sources(self, sources: List[Dict], language: str) -> str:
        """Format sources list"""
        lines = []

        for i, source in enumerate(sources[:15], 1):
            title = source.get("title", "Unknown")
            domain = source.get("domain", "")
            reliability = source.get("reliability_score", 0)
            source_type = source.get("source_type", "web")

            type_labels = {
                "academic": "Akademik" if language == "tr" else "Academic",
                "news": "Haber" if language == "tr" else "News",
                "wiki": "Vikipedi" if language == "tr" else "Wikipedia",
                "web": "Web"
            }

            lines.append(
                f"{i}. {title[:60]}{'...' if len(title) > 60 else ''}\n"
                f"   [{type_labels.get(source_type, 'Web')}] {domain} "
                f"(Güvenilirlik: %{reliability*100:.0f})" if language == "tr" else
                f"   [{type_labels.get(source_type, 'Web')}] {domain} "
                f"(Reliability: {reliability*100:.0f}%)"
            )

        return "\n".join(lines)

    def _extract_keywords(self, findings: List[Dict]) -> List[str]:
        """Extract keywords from findings"""
        all_keywords = []
        for finding in findings:
            keywords = finding.get("keywords", [])
            all_keywords.extend(keywords)

        # Count and return top keywords
        from collections import Counter
        counter = Counter(all_keywords)
        return [word for word, _ in counter.most_common(10)]

    def _generate_docx(
        self,
        metadata: DocumentMetadata,
        sections: List[DocumentSection],
        include_toc: bool
    ) -> Dict[str, Any]:
        """Generate Word document"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE

            doc = Document()

            # Cover Page
            doc.add_spacer = lambda: doc.add_paragraph().add_run().add_break()
            
            # Center title vertically-ish
            for _ in range(5): doc.add_paragraph()
            
            title_para = doc.add_heading(metadata.title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            for _ in range(2): doc.add_paragraph()
            
            # Metadata info on cover
            info_para = doc.add_paragraph()
            info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            info_run = info_para.add_run(
                f"Hazırlayan: {metadata.author}\n"
                f"Tarih: {metadata.date}\n"
                f"Sürüm: {metadata.version}"
            )
            info_run.font.size = Pt(12)
            info_run.font.color.rgb = RGBColor(100, 100, 100)
            
            doc.add_page_break()

            doc.add_page_break()

            # Table of Contents placeholder
            if include_toc:
                doc.add_heading("İçindekiler", 1)
                for i, section in enumerate(sections, 1):
                    toc_para = doc.add_paragraph(f"{i}. {section.title}")
                    toc_para.paragraph_format.left_indent = Inches(0.25)

                    for j, subsection in enumerate(section.subsections, 1):
                        sub_para = doc.add_paragraph(f"   {i}.{j} {subsection.title}")
                        sub_para.paragraph_format.left_indent = Inches(0.5)

                doc.add_page_break()

            # Add sections
            for section in sections:
                doc.add_heading(section.title, section.level)

                if section.content:
                    # Handle bullet points
                    if section.content.startswith("•"):
                        for line in section.content.split("\n"):
                            line = line.strip()
                            if line.startswith("•"):
                                doc.add_paragraph(line[1:].strip(), style='List Bullet')
                            elif line:
                                doc.add_paragraph(line)
                    else:
                        doc.add_paragraph(section.content)

                # Add subsections
                for subsection in section.subsections:
                    doc.add_heading(subsection.title, subsection.level)
                    if subsection.content:
                        for line in subsection.content.split("\n"):
                            line = line.strip()
                            if line.startswith("•"):
                                doc.add_paragraph(line[1:].strip(), style='List Bullet')
                            elif line:
                                doc.add_paragraph(line)

            # Save
            filename = self._sanitize_filename(metadata.title) + ".docx"
            filepath = self.output_dir / filename
            doc.save(str(filepath))

            return {
                "success": True,
                "format": "docx",
                "path": str(filepath),
                "filename": filename,
                "title": metadata.title,
                "sections": len(sections),
                "message": f"Belge oluşturuldu: {filename}"
            }

        except ImportError:
            logger.error("python-docx not installed")
            return self._generate_txt(metadata, sections)
        except Exception as e:
            logger.error(f"DOCX generation error: {e}")
            return {"success": False, "error": str(e)}

    def _generate_pdf(
        self,
        metadata: DocumentMetadata,
        sections: List[DocumentSection],
        include_toc: bool
    ) -> Dict[str, Any]:
        """Generate PDF document"""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
            from reportlab.platypus import Table, TableStyle

            filename = self._sanitize_filename(metadata.title) + ".pdf"
            filepath = self.output_dir / filename

            doc = SimpleDocTemplate(
                str(filepath),
                pagesize=A4,
                rightMargin=72,
                leftMargin=72,
                topMargin=72,
                bottomMargin=72
            )

            styles = getSampleStyleSheet()

            # Custom styles
            styles.add(ParagraphStyle(
                name='CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1  # Center
            ))

            styles.add(ParagraphStyle(
                name='CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                spaceBefore=20,
                spaceAfter=10
            ))

            story = []

            # Title Page
            story.append(Spacer(1, 2*inch))
            story.append(Paragraph(metadata.title, styles['CustomTitle']))
            story.append(Spacer(1, 0.5*inch))
            
            # Metadata
            meta_text = f"""
            <para align="center">
            <font size="12" color="#666666">
            Hazırlayan: {metadata.author}<br/>
            Tarih: {metadata.date}<br/>
            Sürüm: {metadata.version}
            </font>
            </para>
            """
            story.append(Paragraph(meta_text, styles['Normal']))
            story.append(PageBreak())

            # Sections
            for section in sections:
                story.append(Paragraph(section.title, styles['CustomHeading']))

                if section.content:
                    # Split into paragraphs
                    for para in section.content.split("\n\n"):
                        para = para.strip()
                        if para:
                            # Handle bullet points
                            if para.startswith("•"):
                                para = para.replace("•", "&#8226;")
                            story.append(Paragraph(para, styles['Normal']))
                            story.append(Spacer(1, 6))

                # Subsections
                for subsection in section.subsections:
                    story.append(Paragraph(subsection.title, styles['Heading3']))
                    if subsection.content:
                        for para in subsection.content.split("\n\n"):
                            para = para.strip()
                            if para:
                                if para.startswith("•"):
                                    para = para.replace("•", "&#8226;")
                                story.append(Paragraph(para, styles['Normal']))
                                story.append(Spacer(1, 6))

                story.append(Spacer(1, 12))

            doc.build(story)

            return {
                "success": True,
                "format": "pdf",
                "path": str(filepath),
                "filename": filename,
                "title": metadata.title,
                "sections": len(sections),
                "message": f"PDF oluşturuldu: {filename}"
            }

        except ImportError:
            logger.warning("reportlab not installed, falling back to DOCX")
            return self._generate_docx(metadata, sections, include_toc)
        except Exception as e:
            logger.error(f"PDF generation error: {e}")
            return {"success": False, "error": str(e)}

    def _generate_html(
        self,
        metadata: DocumentMetadata,
        sections: List[DocumentSection],
        include_toc: bool
    ) -> Dict[str, Any]:
        """Generate HTML document"""
        try:
            html_parts = [
                "<!DOCTYPE html>",
                "<html lang='tr'>",
                "<head>",
                f"<meta charset='UTF-8'>",
                f"<title>{metadata.title}</title>",
                "<style>",
                """
                body { font-family: 'Segoe UI', Arial, sans-serif; max-width: 900px; margin: 0 auto; padding: 40px; line-height: 1.6; }
                h1 { color: #2c3e50; border-bottom: 3px solid #3498db; padding-bottom: 10px; }
                h2 { color: #34495e; margin-top: 30px; }
                h3 { color: #7f8c8d; }
                .metadata { color: #95a5a6; text-align: center; margin-bottom: 30px; }
                .toc { background: #f8f9fa; padding: 20px; border-radius: 8px; margin-bottom: 30px; }
                .toc a { color: #3498db; text-decoration: none; }
                .toc a:hover { text-decoration: underline; }
                ul { padding-left: 20px; }
                li { margin: 8px 0; }
                .section { margin: 30px 0; }
                .stats { background: #e8f4f8; padding: 15px; border-radius: 5px; }
                blockquote { border-left: 4px solid #3498db; padding-left: 20px; color: #666; }
                """,
                "</style>",
                "</head>",
                "<body>"
            ]

            # Title and metadata
            html_parts.append(f"<h1>{metadata.title}</h1>")
            html_parts.append(f"<div class='metadata'>")
            html_parts.append(f"Hazırlayan: {metadata.author} | Tarih: {metadata.date} | Versiyon: {metadata.version}")
            html_parts.append("</div>")

            # TOC
            if include_toc and sections:
                html_parts.append("<div class='toc'>")
                html_parts.append("<h2>İçindekiler</h2>")
                html_parts.append("<ul>")
                for i, section in enumerate(sections, 1):
                    section_id = f"section-{i}"
                    html_parts.append(f"<li><a href='#{section_id}'>{section.title}</a></li>")
                html_parts.append("</ul>")
                html_parts.append("</div>")

            # Sections
            for i, section in enumerate(sections, 1):
                section_id = f"section-{i}"
                html_parts.append(f"<div class='section' id='{section_id}'>")
                html_parts.append(f"<h2>{section.title}</h2>")

                if section.content:
                    content = section.content.replace("\n", "<br>")
                    content = content.replace("• ", "<li>").replace("<br><li>", "</li><li>")
                    if "<li>" in content:
                        content = "<ul>" + content + "</li></ul>"
                    html_parts.append(f"<p>{content}</p>")

                for subsection in section.subsections:
                    html_parts.append(f"<h3>{subsection.title}</h3>")
                    if subsection.content:
                        sub_content = subsection.content.replace("\n", "<br>")
                        html_parts.append(f"<p>{sub_content}</p>")

                html_parts.append("</div>")

            html_parts.append("</body></html>")

            # Save
            filename = self._sanitize_filename(metadata.title) + ".html"
            filepath = self.output_dir / filename
            filepath.write_text("\n".join(html_parts), encoding="utf-8")

            return {
                "success": True,
                "format": "html",
                "path": str(filepath),
                "filename": filename,
                "title": metadata.title,
                "sections": len(sections),
                "message": f"HTML oluşturuldu: {filename}"
            }

        except Exception as e:
            logger.error(f"HTML generation error: {e}")
            return {"success": False, "error": str(e)}

    def _generate_markdown(
        self,
        metadata: DocumentMetadata,
        sections: List[DocumentSection]
    ) -> Dict[str, Any]:
        """Generate Markdown document"""
        try:
            md_parts = [
                f"# {metadata.title}",
                "",
                f"**Hazırlayan:** {metadata.author}  ",
                f"**Tarih:** {metadata.date}  ",
                f"**Versiyon:** {metadata.version}",
                "",
                "---",
                ""
            ]

            # TOC
            md_parts.append("## İçindekiler")
            md_parts.append("")
            for i, section in enumerate(sections, 1):
                anchor = section.title.lower().replace(" ", "-").replace("ı", "i").replace("ğ", "g")
                md_parts.append(f"{i}. [{section.title}](#{anchor})")
            md_parts.append("")
            md_parts.append("---")
            md_parts.append("")

            # Sections
            for section in sections:
                md_parts.append(f"## {section.title}")
                md_parts.append("")

                if section.content:
                    md_parts.append(section.content)
                    md_parts.append("")

                for subsection in section.subsections:
                    md_parts.append(f"### {subsection.title}")
                    md_parts.append("")
                    if subsection.content:
                        md_parts.append(subsection.content)
                        md_parts.append("")

            # Save
            filename = self._sanitize_filename(metadata.title) + ".md"
            filepath = self.output_dir / filename
            filepath.write_text("\n".join(md_parts), encoding="utf-8")

            return {
                "success": True,
                "format": "markdown",
                "path": str(filepath),
                "filename": filename,
                "title": metadata.title,
                "sections": len(sections),
                "message": f"Markdown oluşturuldu: {filename}"
            }

        except Exception as e:
            logger.error(f"Markdown generation error: {e}")
            return {"success": False, "error": str(e)}

    def _generate_txt(
        self,
        metadata: DocumentMetadata,
        sections: List[DocumentSection]
    ) -> Dict[str, Any]:
        """Generate plain text document"""
        try:
            lines = [
                "=" * 60,
                metadata.title.upper(),
                "=" * 60,
                "",
                f"Hazırlayan: {metadata.author}",
                f"Tarih: {metadata.date}",
                f"Versiyon: {metadata.version}",
                "",
                "-" * 60,
                ""
            ]

            for section in sections:
                lines.append(section.title.upper())
                lines.append("-" * len(section.title))
                lines.append("")

                if section.content:
                    lines.append(section.content)
                    lines.append("")

                for subsection in section.subsections:
                    lines.append(f"  {subsection.title}")
                    lines.append(f"  {'-' * len(subsection.title)}")
                    if subsection.content:
                        for line in subsection.content.split("\n"):
                            lines.append(f"    {line}")
                    lines.append("")

                lines.append("")

            # Save
            filename = self._sanitize_filename(metadata.title) + ".txt"
            filepath = self.output_dir / filename
            filepath.write_text("\n".join(lines), encoding="utf-8")

            return {
                "success": True,
                "format": "txt",
                "path": str(filepath),
                "filename": filename,
                "title": metadata.title,
                "sections": len(sections),
                "message": f"Metin dosyası oluşturuldu: {filename}"
            }

        except Exception as e:
            logger.error(f"TXT generation error: {e}")
            return {"success": False, "error": str(e)}

    def _sanitize_filename(self, name: str) -> str:
        """Sanitize filename"""
        # Remove invalid characters
        invalid_chars = '<>:"/\\|?*'
        for char in invalid_chars:
            name = name.replace(char, '')

        # Replace Turkish characters
        tr_map = {'ı': 'i', 'ğ': 'g', 'ü': 'u', 'ş': 's', 'ö': 'o', 'ç': 'c',
                  'İ': 'I', 'Ğ': 'G', 'Ü': 'U', 'Ş': 'S', 'Ö': 'O', 'Ç': 'C'}
        for tr, en in tr_map.items():
            name = name.replace(tr, en)

        # Limit length
        name = name[:100]

        # Replace spaces with underscores
        name = name.replace(' ', '_')

        return name


# Singleton instance
_generator_instance = None


def get_document_generator() -> ProfessionalDocumentGenerator:
    """Get or create document generator instance"""
    global _generator_instance
    if _generator_instance is None:
        _generator_instance = ProfessionalDocumentGenerator()
    return _generator_instance


async def generate_research_document(
    research_data: Dict[str, Any],
    format: str = "docx",
    template: str = "research_report",
    custom_title: str = None,
    language: str = "tr"
) -> Dict[str, Any]:
    """
    Generate a document from research data

    Args:
        research_data: Research results
        format: Output format (docx/pdf/html/markdown/txt)
        template: Document template
        custom_title: Custom document title
        language: Document language

    Returns:
        Result with file path
    """
    generator = get_document_generator()

    format_map = {
        "docx": DocumentFormat.DOCX,
        "pdf": DocumentFormat.PDF,
        "html": DocumentFormat.HTML,
        "markdown": DocumentFormat.MARKDOWN,
        "md": DocumentFormat.MARKDOWN,
        "txt": DocumentFormat.TXT
    }

    template_map = {
        "research_report": DocumentTemplate.RESEARCH_REPORT,
        "executive_summary": DocumentTemplate.EXECUTIVE_SUMMARY,
        "academic_paper": DocumentTemplate.ACADEMIC_PAPER,
        "business_report": DocumentTemplate.BUSINESS_REPORT
    }

    doc_format = format_map.get(format.lower(), DocumentFormat.DOCX)
    doc_template = template_map.get(template.lower(), DocumentTemplate.RESEARCH_REPORT)

    return generator.generate_research_report(
        research_data=research_data,
        template=doc_template,
        format=doc_format,
        custom_title=custom_title,
        language=language
    )
