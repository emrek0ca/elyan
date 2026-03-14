"""
Word Dosyası Düzenleyici - Word Editor
Replace, add_paragraph, add_heading, format_text ve section-level revizyon operasyonları
"""

import shutil
from pathlib import Path
from datetime import datetime
from typing import Any

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from security.validator import validate_path
from utils.logger import get_logger

logger = get_logger("word_editor")


SECTION_TITLES = {
    "kısa özet": "Kısa Özet",
    "kisa ozet": "Kısa Özet",
    "özet": "Kısa Özet",
    "ozet": "Kısa Özet",
    "temel bulgular": "Temel Bulgular",
    "bulgular": "Temel Bulgular",
    "sonuç": "Sonuç",
    "sonuc": "Sonuç",
    "açık riskler": "Açık Riskler",
    "acik riskler": "Açık Riskler",
    "riskler": "Açık Riskler",
    "belirsizlikler": "Belirsizlikler",
    "kaynakça": "Kaynakça",
    "kaynaklar": "Kaynakça",
}


def _normalize_section_name(value: str) -> str:
    raw = str(value or "").strip().lower().rstrip(":")
    return SECTION_TITLES.get(raw, str(value or "").strip())


def _paragraph_is_section_heading(paragraph) -> bool:
    text = str(paragraph.text or "").strip()
    if not text:
        return False
    style_name = str(getattr(paragraph.style, "name", "") or "").lower()
    if style_name.startswith("heading"):
        return True
    normalized = _normalize_section_name(text)
    return normalized in set(SECTION_TITLES.values())


def _find_section_range(doc, section_name: str) -> tuple[int, int] | None:
    target = _normalize_section_name(section_name)
    if not target:
        return None
    start = None
    end = len(doc.paragraphs)
    for idx, paragraph in enumerate(doc.paragraphs):
        current = _normalize_section_name(paragraph.text)
        if current == target:
            start = idx
            break
    if start is None:
        return None
    for idx in range(start + 1, len(doc.paragraphs)):
        if _paragraph_is_section_heading(doc.paragraphs[idx]):
            end = idx
            break
    return start, end


def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def _insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_para = paragraph.insert_paragraph_before(text)
    paragraph._p.addnext(new_para._p)
    if style:
        try:
            new_para.style = style
        except KeyError:
            pass
    return new_para


def _replace_section_content(
    doc,
    section_name: str,
    paragraphs: list[str],
    *,
    replace_heading: bool = False,
    heading_override: str = "",
) -> dict[str, Any]:
    section_range = _find_section_range(doc, section_name)
    normalized_name = _normalize_section_name(section_name)
    final_heading = str(heading_override or normalized_name).strip() or normalized_name
    if section_range is None:
        heading = doc.add_heading(final_heading, level=1)
        anchor = heading
        created = True
    else:
        start, end = section_range
        heading_para = doc.paragraphs[start]
        if replace_heading:
            heading_para.text = final_heading
        for idx in range(end - 1, start, -1):
            _delete_paragraph(doc.paragraphs[idx])
        anchor = heading_para
        created = False

    inserted = 0
    last_para = anchor
    for item in [str(p or "").strip() for p in paragraphs if str(p or "").strip()]:
        last_para = _insert_paragraph_after(last_para, item)
        inserted += 1
    if inserted == 0:
        last_para = _insert_paragraph_after(last_para, "")
        inserted = 1
    return {"section": final_heading, "created": created, "paragraphs_written": inserted}


def _build_revision_summary(changes: list[dict[str, Any]]) -> str:
    section_changes: list[str] = []
    generic_changes: list[str] = []
    for change in list(changes or []):
        if not isinstance(change, dict):
            continue
        change_type = str(change.get("type") or "").strip()
        section = str(change.get("section") or "").strip()
        if section:
            section_changes.append(f"{section}: {change_type}")
        elif change_type:
            generic_changes.append(change_type)
    lines = ["Revision Summary", ""]
    if section_changes:
        lines.append("Section updates:")
        for item in section_changes:
            lines.append(f"- {item}")
    if generic_changes:
        lines.append("")
        lines.append("Other operations:")
        for item in generic_changes:
            lines.append(f"- {item}")
    if len(lines) == 2:
        lines.append("No notable changes recorded.")
    return "\n".join(lines).strip() + "\n"


async def edit_word_document(
    path: str,
    operations: list[dict],
    create_backup: bool = True
) -> dict[str, Any]:
    """
    Word dosyasını düzenle

    Args:
        path: Dosya yolu (.docx)
        operations: Düzenleme operasyonları listesi
            Her operasyon:
            - {"type": "replace_text", "find": str, "replace": str}
            - {"type": "add_paragraph", "text": str, "style": str, "position": int}
            - {"type": "add_heading", "text": str, "level": int, "position": int}
            - {"type": "format_text", "find": str, "bold": bool, "italic": bool, "underline": bool}
            - {"type": "add_table", "rows": int, "cols": int, "data": list[list]}
            - {"type": "delete_paragraph", "index": int}
            - {"type": "rewrite_section", "section": str, "content": str | list[str]}
            - {"type": "replace_section", "section": str, "content": str | list[str], "heading": str | None}
            - {"type": "append_risk_note", "text": str}
            - {"type": "generate_revision_summary"}
        create_backup: Yedek oluştur

    Returns:
        dict: Düzenleme sonucu
    """
    if not DOCX_AVAILABLE:
        return {"success": False, "error": "python-docx kütüphanesi yüklü değil"}

    try:
        valid, msg, resolved_path = validate_path(path)
        if not valid:
            return {"success": False, "error": msg}

        if not resolved_path.exists():
            return {"success": False, "error": f"Dosya bulunamadı: {path}"}

        if not resolved_path.suffix.lower() in [".docx", ".doc"]:
            return {"success": False, "error": "Sadece .docx ve .doc dosyaları düzenlenebilir"}

        # Check file size
        size = resolved_path.stat().st_size
        if size > 50 * 1024 * 1024:  # 50MB limit
            return {"success": False, "error": "Dosya çok büyük (max 50MB)"}

        # Create backup
        backup_path = None
        if create_backup:
            backup_path = resolved_path.with_suffix(
                f".backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}{resolved_path.suffix}"
            )
            shutil.copy2(resolved_path, backup_path)
            logger.info(f"Yedek oluşturuldu: {backup_path.name}")

        # Open document
        doc = Document(str(resolved_path))
        changes_made = []
        revision_summary_requested = False

        # Apply operations
        for i, op in enumerate(operations):
            op_type = op.get("type", "").lower()

            if op_type == "replace_text":
                find_text = op.get("find", "")
                replace_text = op.get("replace", "")

                if not find_text:
                    changes_made.append({"op": i+1, "type": "replace_text", "error": "find parametresi gerekli"})
                    continue

                count = 0
                # Replace in paragraphs
                for paragraph in doc.paragraphs:
                    if find_text in paragraph.text:
                        for run in paragraph.runs:
                            if find_text in run.text:
                                run.text = run.text.replace(find_text, replace_text)
                                count += 1

                # Replace in tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if find_text in cell.text:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        if find_text in run.text:
                                            run.text = run.text.replace(find_text, replace_text)
                                            count += 1

                changes_made.append({
                    "op": i+1,
                    "type": "replace_text",
                    "find": find_text[:50],
                    "count": count
                })

            elif op_type == "add_paragraph":
                text = op.get("text", "")
                style = op.get("style")
                position = op.get("position")

                if position is not None and 0 <= position < len(doc.paragraphs):
                    # Insert at position
                    new_para = doc.paragraphs[position].insert_paragraph_before(text)
                else:
                    # Add at end
                    new_para = doc.add_paragraph(text)

                if style:
                    try:
                        new_para.style = style
                    except KeyError:
                        pass

                changes_made.append({
                    "op": i+1,
                    "type": "add_paragraph",
                    "text": text[:50],
                    "position": position
                })

            elif op_type == "add_heading":
                text = op.get("text", "")
                level = op.get("level", 1)
                position = op.get("position")

                if not 0 <= level <= 9:
                    level = 1

                if position is not None and 0 <= position < len(doc.paragraphs):
                    # Insert heading at position
                    heading = doc.paragraphs[position].insert_paragraph_before(text)
                    heading.style = f"Heading {level}"
                else:
                    # Add at end
                    doc.add_heading(text, level=level)

                changes_made.append({
                    "op": i+1,
                    "type": "add_heading",
                    "text": text[:50],
                    "level": level,
                    "position": position
                })

            elif op_type == "format_text":
                find_text = op.get("find", "")
                bold = op.get("bold")
                italic = op.get("italic")
                underline = op.get("underline")
                font_size = op.get("font_size")
                color = op.get("color")

                if not find_text:
                    changes_made.append({"op": i+1, "type": "format_text", "error": "find parametresi gerekli"})
                    continue

                count = 0
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        if find_text in run.text:
                            if bold is not None:
                                run.bold = bold
                            if italic is not None:
                                run.italic = italic
                            if underline is not None:
                                run.underline = underline
                            if font_size:
                                run.font.size = Pt(font_size)
                            if color:
                                try:
                                    # Color as hex string "#RRGGBB"
                                    if color.startswith("#"):
                                        color = color[1:]
                                    r = int(color[0:2], 16)
                                    g = int(color[2:4], 16)
                                    b = int(color[4:6], 16)
                                    run.font.color.rgb = RGBColor(r, g, b)
                                except (ValueError, IndexError):
                                    pass
                            count += 1

                changes_made.append({
                    "op": i+1,
                    "type": "format_text",
                    "find": find_text[:50],
                    "count": count
                })

            elif op_type == "add_table":
                rows = op.get("rows", 2)
                cols = op.get("cols", 2)
                data = op.get("data", [])
                headers = op.get("headers", [])

                table = doc.add_table(rows=rows, cols=cols)
                table.style = "Table Grid"

                # Add headers
                if headers:
                    for j, header in enumerate(headers[:cols]):
                        table.cell(0, j).text = str(header)
                        # Make headers bold
                        for para in table.cell(0, j).paragraphs:
                            for run in para.runs:
                                run.bold = True

                # Add data
                start_row = 1 if headers else 0
                for row_idx, row_data in enumerate(data):
                    actual_row = start_row + row_idx
                    if actual_row >= rows:
                        break
                    for col_idx, cell_value in enumerate(row_data):
                        if col_idx >= cols:
                            break
                        table.cell(actual_row, col_idx).text = str(cell_value)

                changes_made.append({
                    "op": i+1,
                    "type": "add_table",
                    "rows": rows,
                    "cols": cols
                })

            elif op_type == "delete_paragraph":
                index = op.get("index", -1)

                if 0 <= index < len(doc.paragraphs):
                    para = doc.paragraphs[index]
                    _delete_paragraph(para)
                    changes_made.append({
                        "op": i+1,
                        "type": "delete_paragraph",
                        "index": index
                    })
                else:
                    changes_made.append({
                        "op": i+1,
                        "type": "delete_paragraph",
                        "error": f"Geçersiz paragraf indeksi: {index}"
                    })

            elif op_type in {"rewrite_section", "replace_section"}:
                section = str(op.get("section") or op.get("heading") or "").strip()
                if not section:
                    changes_made.append({"op": i+1, "type": op_type, "error": "section parametresi gerekli"})
                    continue
                raw_content = op.get("content")
                if isinstance(raw_content, list):
                    section_paragraphs = [str(item or "").strip() for item in raw_content if str(item or "").strip()]
                else:
                    section_paragraphs = [item.strip() for item in str(raw_content or "").split("\n\n") if item.strip()]
                heading = str(op.get("heading") or section).strip() or section
                result = _replace_section_content(
                    doc,
                    section,
                    section_paragraphs,
                    replace_heading=bool(op_type == "replace_section"),
                    heading_override=heading if op_type == "replace_section" else "",
                )
                changes_made.append(
                    {
                        "op": i + 1,
                        "type": op_type,
                        "section": result.get("section"),
                        "created": bool(result.get("created")),
                        "paragraphs_written": int(result.get("paragraphs_written", 0) or 0),
                    }
                )

            elif op_type == "append_risk_note":
                note = str(op.get("text") or "").strip()
                if not note:
                    changes_made.append({"op": i+1, "type": op_type, "error": "text parametresi gerekli"})
                    continue
                section_range = _find_section_range(doc, "Açık Riskler")
                created = False
                if section_range is None:
                    heading = doc.add_heading("Açık Riskler", level=1)
                    anchor = heading
                    created = True
                else:
                    start, end = section_range
                    anchor = doc.paragraphs[end - 1] if end - 1 >= start else doc.paragraphs[start]
                _insert_paragraph_after(anchor, note)
                changes_made.append(
                    {
                        "op": i + 1,
                        "type": op_type,
                        "section": "Açık Riskler",
                        "created": created,
                        "paragraphs_written": 1,
                    }
                )

            elif op_type == "generate_revision_summary":
                revision_summary_requested = True
                changes_made.append({"op": i + 1, "type": op_type, "status": "queued"})

            else:
                changes_made.append({
                    "op": i+1,
                    "type": op_type,
                    "error": f"Bilinmeyen operasyon tipi: {op_type}"
                })

        # Save document
        doc.save(str(resolved_path))
        revision_summary = _build_revision_summary(changes_made) if revision_summary_requested else ""
        artifacts: list[str] = [str(resolved_path)]
        revision_summary_path = None
        if revision_summary_requested:
            revision_summary_path = resolved_path.with_suffix(".revision_summary.md")
            revision_summary_path.write_text(revision_summary, encoding="utf-8")
            artifacts.append(str(revision_summary_path))

        logger.info(f"Word dosyası düzenlendi: {resolved_path.name} - {len(changes_made)} operasyon")

        return {
            "success": True,
            "path": str(resolved_path),
            "filename": resolved_path.name,
            "changes": changes_made,
            "backup_created": create_backup,
            "backup_path": str(backup_path) if backup_path else "",
            "revision_summary": revision_summary,
            "revision_summary_path": str(revision_summary_path) if revision_summary_path else "",
            "artifacts": artifacts,
            "message": f"Word dosyası düzenlendi: {len(changes_made)} operasyon uygulandı"
        }

    except Exception as e:
        logger.error(f"Word düzenleme hatası: {e}")
        return {"success": False, "error": f"Word dosyası düzenlenemedi: {str(e)}"}


async def get_word_structure(path: str) -> dict[str, Any]:
    """
    Word dosyasının yapısını getir

    Args:
        path: Dosya yolu

    Returns:
        dict: Belge yapısı (paragraflar, tablolar, vs.)
    """
    if not DOCX_AVAILABLE:
        return {"success": False, "error": "python-docx kütüphanesi yüklü değil"}

    try:
        valid, msg, resolved_path = validate_path(path)
        if not valid:
            return {"success": False, "error": msg}

        if not resolved_path.exists():
            return {"success": False, "error": f"Dosya bulunamadı: {path}"}

        doc = Document(str(resolved_path))

        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            paragraphs.append({
                "index": i,
                "style": para.style.name if para.style else None,
                "text": para.text[:100] + "..." if len(para.text) > 100 else para.text
            })

        tables = []
        for i, table in enumerate(doc.tables):
            tables.append({
                "index": i,
                "rows": len(table.rows),
                "cols": len(table.columns)
            })

        return {
            "success": True,
            "path": str(resolved_path),
            "filename": resolved_path.name,
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
            "paragraphs": paragraphs[:50],  # Limit to first 50
            "tables": tables
        }

    except Exception as e:
        logger.error(f"Word yapı okuma hatası: {e}")
        return {"success": False, "error": f"Word yapısı okunamadı: {str(e)}"}
