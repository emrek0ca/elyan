"""
Belge Birleştirici - Document Merger
PDF ve Word dosyalarını birleştirme
"""

from pathlib import Path
from datetime import datetime
from typing import Any

try:
    from pypdf import PdfReader, PdfWriter
    PYPDF_AVAILABLE = True
except ImportError:
    try:
        from PyPDF2 import PdfReader, PdfWriter
        PYPDF_AVAILABLE = True
    except ImportError:
        PYPDF_AVAILABLE = False

try:
    from docx import Document
    from docxcompose.composer import Composer
    DOCX_AVAILABLE = True
    DOCXCOMPOSE_AVAILABLE = True
except ImportError:
    try:
        from docx import Document
        DOCX_AVAILABLE = True
        DOCXCOMPOSE_AVAILABLE = False
    except ImportError:
        DOCX_AVAILABLE = False
        DOCXCOMPOSE_AVAILABLE = False

from security.validator import validate_path
from config.settings import HOME_DIR
from utils.logger import get_logger

logger = get_logger("document_merger")

# Maximum files to merge
MAX_MERGE_FILES = 10


async def merge_documents(
    input_paths: list[str],
    output_path: str,
    output_format: str = "auto"
) -> dict[str, Any]:
    """
    Belgeleri birleştir (PDF veya Word)

    Args:
        input_paths: Birleştirilecek dosya yolları
        output_path: Çıktı dosya yolu
        output_format: Çıktı formatı ("auto", "pdf", "docx")

    Returns:
        dict: Birleştirme sonucu
    """
    try:
        if not input_paths:
            return {"success": False, "error": "En az bir dosya gerekli"}

        if len(input_paths) > MAX_MERGE_FILES:
            return {
                "success": False,
                "error": f"Maksimum {MAX_MERGE_FILES} dosya birleştirilebilir"
            }

        # Validate all input paths
        validated_paths = []
        for path in input_paths:
            valid, msg, resolved = validate_path(path)
            if not valid:
                return {"success": False, "error": f"Geçersiz yol: {path} - {msg}"}
            if not resolved.exists():
                return {"success": False, "error": f"Dosya bulunamadı: {path}"}
            validated_paths.append(resolved)

        # Validate output path
        valid, msg, output_resolved = validate_path(output_path)
        if not valid:
            return {"success": False, "error": f"Geçersiz çıktı yolu: {msg}"}

        # Ensure output directory exists
        output_resolved.parent.mkdir(parents=True, exist_ok=True)

        # Determine format
        if output_format == "auto":
            first_ext = validated_paths[0].suffix.lower()
            if first_ext == ".pdf":
                output_format = "pdf"
            elif first_ext in [".docx", ".doc"]:
                output_format = "docx"
            else:
                return {
                    "success": False,
                    "error": f"Bilinmeyen dosya formatı: {first_ext}"
                }

        # Check all files have same format
        for path in validated_paths:
            ext = path.suffix.lower()
            if output_format == "pdf" and ext != ".pdf":
                return {
                    "success": False,
                    "error": f"Tüm dosyalar aynı formatta olmalı. {path.name} PDF değil."
                }
            if output_format == "docx" and ext not in [".docx", ".doc"]:
                return {
                    "success": False,
                    "error": f"Tüm dosyalar aynı formatta olmalı. {path.name} Word değil."
                }

        # Merge based on format
        if output_format == "pdf":
            return await merge_pdfs(
                [str(p) for p in validated_paths],
                str(output_resolved)
            )
        elif output_format == "docx":
            return await merge_word_documents(
                [str(p) for p in validated_paths],
                str(output_resolved)
            )
        else:
            return {"success": False, "error": f"Desteklenmeyen format: {output_format}"}

    except Exception as e:
        logger.error(f"Belge birleştirme hatası: {e}")
        return {"success": False, "error": f"Belgeler birleştirilemedi: {str(e)}"}


async def merge_pdfs(
    input_paths: list[str],
    output_path: str,
    page_ranges: dict[str, str] | None = None
) -> dict[str, Any]:
    """
    PDF dosyalarını birleştir

    Args:
        input_paths: Birleştirilecek PDF dosya yolları
        output_path: Çıktı PDF dosya yolu
        page_ranges: Dosya başına sayfa aralıkları
            Örnek: {"rapor1.pdf": "1-5", "rapor2.pdf": "2,4,6-10"}

    Returns:
        dict: Birleştirme sonucu
    """
    if not PYPDF_AVAILABLE:
        return {"success": False, "error": "pypdf veya PyPDF2 kütüphanesi yüklü değil"}

    try:
        if not input_paths:
            return {"success": False, "error": "En az bir PDF dosyası gerekli"}

        if len(input_paths) > MAX_MERGE_FILES:
            return {
                "success": False,
                "error": f"Maksimum {MAX_MERGE_FILES} dosya birleştirilebilir"
            }

        # Validate paths
        validated_paths = []
        for path in input_paths:
            valid, msg, resolved = validate_path(path)
            if not valid:
                return {"success": False, "error": f"Geçersiz yol: {path} - {msg}"}
            if not resolved.exists():
                return {"success": False, "error": f"Dosya bulunamadı: {path}"}
            if resolved.suffix.lower() != ".pdf":
                return {"success": False, "error": f"PDF dosyası değil: {path}"}
            validated_paths.append(resolved)

        # Validate output path
        valid, msg, output_resolved = validate_path(output_path)
        if not valid:
            return {"success": False, "error": f"Geçersiz çıktı yolu: {msg}"}

        # Ensure .pdf extension
        if output_resolved.suffix.lower() != ".pdf":
            output_resolved = output_resolved.with_suffix(".pdf")

        output_resolved.parent.mkdir(parents=True, exist_ok=True)

        # Create PDF writer
        writer = PdfWriter()
        total_pages = 0
        file_info = []

        for path in validated_paths:
            try:
                reader = PdfReader(str(path))
                num_pages = len(reader.pages)

                # Determine which pages to include
                if page_ranges and path.name in page_ranges:
                    pages_to_add = _parse_page_range(page_ranges[path.name], num_pages)
                else:
                    pages_to_add = list(range(num_pages))

                # Add pages
                added_count = 0
                for page_idx in pages_to_add:
                    if 0 <= page_idx < num_pages:
                        writer.add_page(reader.pages[page_idx])
                        added_count += 1

                total_pages += added_count
                file_info.append({
                    "filename": path.name,
                    "total_pages": num_pages,
                    "pages_added": added_count
                })

            except Exception as e:
                logger.warning(f"PDF okuma hatası ({path.name}): {e}")
                file_info.append({
                    "filename": path.name,
                    "error": str(e)
                })

        if total_pages == 0:
            return {"success": False, "error": "Birleştirilecek sayfa bulunamadı"}

        # Write output
        with open(output_resolved, "wb") as f:
            writer.write(f)

        logger.info(f"PDF birleştirildi: {output_resolved.name} - {total_pages} sayfa")

        return {
            "success": True,
            "output_path": str(output_resolved),
            "filename": output_resolved.name,
            "total_pages": total_pages,
            "input_files": len(validated_paths),
            "file_info": file_info,
            "message": f"PDF birleştirildi: {len(validated_paths)} dosya, {total_pages} sayfa"
        }

    except Exception as e:
        logger.error(f"PDF birleştirme hatası: {e}")
        return {"success": False, "error": f"PDF dosyaları birleştirilemedi: {str(e)}"}


async def merge_word_documents(
    input_paths: list[str],
    output_path: str
) -> dict[str, Any]:
    """
    Word dosyalarını birleştir

    Args:
        input_paths: Birleştirilecek Word dosya yolları
        output_path: Çıktı Word dosya yolu

    Returns:
        dict: Birleştirme sonucu
    """
    if not DOCX_AVAILABLE:
        return {"success": False, "error": "python-docx kütüphanesi yüklü değil"}

    try:
        if not input_paths:
            return {"success": False, "error": "En az bir Word dosyası gerekli"}

        if len(input_paths) > MAX_MERGE_FILES:
            return {
                "success": False,
                "error": f"Maksimum {MAX_MERGE_FILES} dosya birleştirilebilir"
            }

        # Validate paths
        validated_paths = []
        for path in input_paths:
            valid, msg, resolved = validate_path(path)
            if not valid:
                return {"success": False, "error": f"Geçersiz yol: {path} - {msg}"}
            if not resolved.exists():
                return {"success": False, "error": f"Dosya bulunamadı: {path}"}
            if resolved.suffix.lower() not in [".docx", ".doc"]:
                return {"success": False, "error": f"Word dosyası değil: {path}"}
            validated_paths.append(resolved)

        # Validate output path
        valid, msg, output_resolved = validate_path(output_path)
        if not valid:
            return {"success": False, "error": f"Geçersiz çıktı yolu: {msg}"}

        # Ensure .docx extension
        if output_resolved.suffix.lower() not in [".docx", ".doc"]:
            output_resolved = output_resolved.with_suffix(".docx")

        output_resolved.parent.mkdir(parents=True, exist_ok=True)

        file_info = []

        if DOCXCOMPOSE_AVAILABLE:
            # Use docxcompose for better merging
            try:
                master = Document(str(validated_paths[0]))
                composer = Composer(master)

                file_info.append({
                    "filename": validated_paths[0].name,
                    "paragraphs": len(master.paragraphs)
                })

                for path in validated_paths[1:]:
                    doc = Document(str(path))
                    composer.append(doc)
                    file_info.append({
                        "filename": path.name,
                        "paragraphs": len(doc.paragraphs)
                    })

                composer.save(str(output_resolved))

            except Exception as e:
                logger.warning(f"docxcompose hatası, basit birleştirmeye geçiliyor: {e}")
                # Fallback to simple merge
                return await _simple_word_merge(validated_paths, output_resolved)

        else:
            # Simple merge without docxcompose
            return await _simple_word_merge(validated_paths, output_resolved)

        logger.info(f"Word birleştirildi: {output_resolved.name} - {len(validated_paths)} dosya")

        return {
            "success": True,
            "output_path": str(output_resolved),
            "filename": output_resolved.name,
            "input_files": len(validated_paths),
            "file_info": file_info,
            "message": f"Word dosyaları birleştirildi: {len(validated_paths)} dosya"
        }

    except Exception as e:
        logger.error(f"Word birleştirme hatası: {e}")
        return {"success": False, "error": f"Word dosyaları birleştirilemedi: {str(e)}"}


async def _simple_word_merge(
    input_paths: list[Path],
    output_path: Path
) -> dict[str, Any]:
    """
    Basit Word birleştirme (docxcompose olmadan)
    """
    try:
        # Create new document
        merged_doc = Document()
        file_info = []
        total_paragraphs = 0

        for i, path in enumerate(input_paths):
            doc = Document(str(path))

            # Add page break between documents (except first)
            if i > 0:
                merged_doc.add_page_break()

            # Add document title as heading
            merged_doc.add_heading(path.stem, level=1)

            # Copy paragraphs
            para_count = 0
            for para in doc.paragraphs:
                new_para = merged_doc.add_paragraph(para.text)
                # Try to preserve style
                if para.style:
                    try:
                        new_para.style = para.style.name
                    except:
                        pass
                para_count += 1

            # Copy tables
            for table in doc.tables:
                # Create new table
                new_table = merged_doc.add_table(
                    rows=len(table.rows),
                    cols=len(table.columns)
                )
                new_table.style = "Table Grid"

                for row_idx, row in enumerate(table.rows):
                    for col_idx, cell in enumerate(row.cells):
                        new_table.cell(row_idx, col_idx).text = cell.text

            total_paragraphs += para_count
            file_info.append({
                "filename": path.name,
                "paragraphs": para_count,
                "tables": len(doc.tables)
            })

        merged_doc.save(str(output_path))

        return {
            "success": True,
            "output_path": str(output_path),
            "filename": output_path.name,
            "input_files": len(input_paths),
            "total_paragraphs": total_paragraphs,
            "file_info": file_info,
            "message": f"Word dosyaları birleştirildi: {len(input_paths)} dosya"
        }

    except Exception as e:
        logger.error(f"Basit Word birleştirme hatası: {e}")
        return {"success": False, "error": f"Word dosyaları birleştirilemedi: {str(e)}"}


def _parse_page_range(range_str: str, max_pages: int) -> list[int]:
    """
    Sayfa aralığı string'ini sayfa indekslerine çevir

    Args:
        range_str: Sayfa aralığı (örn: "1-5,7,9-12")
        max_pages: Maksimum sayfa sayısı

    Returns:
        list[int]: Sayfa indeksleri (0-indexed)
    """
    pages = set()
    parts = range_str.split(",")

    for part in parts:
        part = part.strip()
        if "-" in part:
            try:
                start, end = part.split("-", 1)
                start = int(start.strip()) - 1  # Convert to 0-indexed
                end = int(end.strip())  # Keep end as 1-indexed for range
                for i in range(max(0, start), min(end, max_pages)):
                    pages.add(i)
            except ValueError:
                continue
        else:
            try:
                page = int(part) - 1  # Convert to 0-indexed
                if 0 <= page < max_pages:
                    pages.add(page)
            except ValueError:
                continue

    return sorted(pages)
