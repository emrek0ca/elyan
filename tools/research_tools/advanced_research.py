"""
Gelişmiş Araştırma - Advanced Research
Çoklu kaynak araştırma, derinlik seviyeleri ve kaynak değerlendirmesi
"""

import asyncio
import re
import uuid
from collections import Counter
from datetime import datetime
from pathlib import Path
from typing import Any
from enum import Enum
from dataclasses import dataclass, field
from urllib.parse import urlparse
from utils.logger import get_logger

logger = get_logger("advanced_research")

# Global cache for last research topic (used by report generation)
_last_research_topic = None
_last_research_result: dict[str, Any] | None = None
RESEARCH_REPORT_DIR = Path.home() / ".elyan" / "reports" / "research"

_TRUSTED_DOMAIN_HINTS = (
    "wikipedia.org",
    "who.int",
    "cdc.gov",
    "nih.gov",
    "mayoclinic.org",
    "akc.org",
    "vca",
    "reuters.com",
    "bbc.",
    "apnews.com",
)
_ACADEMIC_DOMAIN_HINTS = (
    ".edu",
    ".ac.",
    "arxiv.org",
    "scholar.google",
    "researchgate.net",
    "academia.edu",
    "springer.com",
    "sciencedirect.com",
    "nature.com",
)
_OFFICIAL_DOMAIN_HINTS = (
    ".gov",
    "who.int",
    "cdc.gov",
    "nih.gov",
    "europa.eu",
)
_LOW_VALUE_DOMAIN_HINTS = (
    "blogspot.",
    "wordpress.",
    "tumblr.",
    "pinterest.",
    "tiktok.",
    "instagram.com",
    "facebook.com",
    "x.com",
    "twitter.com",
    "reddit.com",
)
_LOW_VALUE_URL_PARTS = (
    "/tag/",
    "/tags/",
    "/kategori/",
    "/category/",
    "/search",
    "/etiket/",
    "/amp",
    "?amp",
)
_NOISE_PATTERNS = (
    "cookie",
    "çerez",
    "gdpr",
    "privacy policy",
    "gizlilik",
    "kampanya",
    "indirim",
    "satın al",
    "hemen al",
    "newsletter",
    "abone ol",
    "devamını oku",
    "yorum yap",
    "paylaş",
    "share",
    "oturum aç",
    "giriş yap",
    "login",
)
_STOPWORDS = {
    "ve", "ile", "bir", "bu", "için", "olan", "de", "da", "the", "and", "is", "of", "to", "in", "a",
    "veya", "daha", "çok", "gibi", "ama", "fakat", "hem", "güncel", "genel", "üzerine", "hakkında",
}

_SOURCE_POLICIES = {"balanced", "trusted", "academic", "official"}


def _normalize_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _domain_from_url(url: str) -> str:
    try:
        domain = urlparse(str(url or "")).netloc.lower().strip()
        if domain.startswith("www."):
            domain = domain[4:]
        return domain
    except Exception:
        return ""


def _normalize_source_policy(policy: str) -> str:
    raw = _normalize_text(policy).lower().replace("-", "_")
    alias_map = {
        "safe": "trusted",
        "reliable": "trusted",
        "secure": "trusted",
        "akademik": "academic",
        "bilimsel": "academic",
        "academic_only": "academic",
        "resmi": "official",
        "official_only": "official",
        "government": "official",
    }
    if raw in alias_map:
        raw = alias_map[raw]
    return raw if raw in _SOURCE_POLICIES else "balanced"


def _matches_policy_domain(domain: str, policy: str) -> bool:
    d = str(domain or "").lower().strip()
    if not d:
        return False

    if policy == "balanced":
        return not any(h in d for h in _LOW_VALUE_DOMAIN_HINTS)

    if policy == "trusted":
        if any(h in d for h in _LOW_VALUE_DOMAIN_HINTS):
            return False
        if d.endswith((".gov", ".edu")):
            return True
        if any(h in d for h in _TRUSTED_DOMAIN_HINTS):
            return True
        if d.endswith(".org"):
            return True
        return False

    if policy == "academic":
        if d.endswith((".edu", ".ac.uk", ".ac.jp", ".ac.tr")):
            return True
        return any(h in d for h in _ACADEMIC_DOMAIN_HINTS)

    if policy == "official":
        if d.endswith(".gov"):
            return True
        return any(h in d for h in _OFFICIAL_DOMAIN_HINTS)

    return True


def _apply_source_policy(results: list[dict[str, Any]], policy: str, target_sources: int) -> list[dict[str, Any]]:
    norm_policy = _normalize_source_policy(policy)
    if not results:
        return []
    if norm_policy == "balanced":
        # Balanced mode already handled by quality scoring; keep ordering intact.
        return list(results[: max(1, int(target_sources))])

    filtered: list[dict[str, Any]] = []
    fallback: list[dict[str, Any]] = []
    for item in results:
        if not isinstance(item, dict):
            continue
        domain = _domain_from_url(item.get("url", ""))
        if _matches_policy_domain(domain, norm_policy):
            filtered.append(item)
        else:
            fallback.append(item)

    need = max(1, int(target_sources))
    if len(filtered) >= need:
        return filtered[:need]

    # If strict policy yields too few results, softly backfill from balanced pool.
    merged = list(filtered)
    for item in fallback:
        merged.append(item)
        if len(merged) >= need:
            break
    return merged[:need]


def _source_policy_stats(results: list[dict[str, Any]], policy: str) -> dict[str, Any]:
    norm = _normalize_source_policy(policy)
    if not results:
        return {"policy": norm, "compliance_ratio": 0.0, "fallback_used": False}
    if norm == "balanced":
        return {"policy": norm, "compliance_ratio": 1.0, "fallback_used": False}
    total = len(results)
    matched = 0
    for item in results:
        if not isinstance(item, dict):
            continue
        domain = _domain_from_url(item.get("url", ""))
        if _matches_policy_domain(domain, norm):
            matched += 1
    ratio = matched / max(total, 1)
    return {
        "policy": norm,
        "compliance_ratio": round(ratio, 2),
        "fallback_used": bool(ratio < 1.0),
        "matched": matched,
        "total": total,
    }


def _apply_min_reliability(sources: list["ResearchSource"], min_reliability: float, keep_at_least: int = 2) -> list["ResearchSource"]:
    if not sources:
        return []
    threshold = max(0.0, min(1.0, float(min_reliability or 0.0)))
    scored = sorted(sources, key=lambda s: float(s.reliability_score or 0.0), reverse=True)
    strong = [s for s in scored if float(s.reliability_score or 0.0) >= threshold]
    if len(strong) >= keep_at_least:
        return strong
    # Fallback to top-N instead of returning almost empty result.
    return scored[: max(1, keep_at_least)]


def _tokenize_topic(topic: str) -> list[str]:
    tokens = []
    for tok in re.findall(r"[a-zA-ZğüşöçıİĞÜŞÖÇ0-9]+", _normalize_text(topic).lower()):
        if len(tok) >= 3 and tok not in _STOPWORDS:
            tokens.append(tok)
    return tokens


def _is_noise_sentence(text: str) -> bool:
    t = _normalize_text(text).lower()
    if not t:
        return True
    if len(t) < 55:
        return True
    if len(t) > 320:
        return True
    if t.endswith("?"):
        return True
    if not re.search(r"[.!]$", t) and len(t) > 170:
        return True
    if t.count(",") >= 8:
        return True
    if re.search(r"\b\d+\s*(?:milyon\w*|milyar\w*)\b.*\b(?:ırk|irk|cins|tür|tur)\w*\b", t):
        return True
    # Menu/navigation-like keyword dumps are common in pet/blog pages.
    if re.search(r"\b(?:toy|poodle|maltese|maltipoo|yorkshire|pomeranian|bulldog|retriever|terrier|chow|cavalier)\b", t):
        breed_hits = len(re.findall(r"\b(?:toy|poodle|maltese|maltipoo|yorkshire|pomeranian|bulldog|retriever|terrier|chow|cavalier)\b", t))
        if breed_hits >= 4 and not re.search(r"[.!?]", t):
            return True
    # Repeated short title fragments without a verbal predicate are usually nav blocks.
    if t.count(" ") >= 12 and not re.search(r"[.!?]", t):
        if len(re.findall(r"\b(?:en|ile|ve|için|icin|listesi|cinsleri|türleri|turleri)\b", t)) >= 5:
            return True
    if any(p in t for p in _NOISE_PATTERNS):
        return True
    words = t.split()
    if len(words) > 45:
        return True
    return False


def _compact_finding_text(text: str, max_len: int = 210) -> str:
    t = _normalize_text(text)
    if len(t) <= max_len:
        return t
    cropped = t[:max_len].rstrip(" ,;:-")
    for sep in (". ", "; ", " - ", ", "):
        idx = cropped.rfind(sep)
        if idx >= 90:
            cropped = cropped[:idx + 1].rstrip(" ,;:-")
            break
    if not re.search(r"[.!?]$", cropped):
        cropped += "."
    return cropped


def _split_sentences(text: str) -> list[str]:
    raw = _normalize_text(text)
    if not raw:
        return []
    # Keep punctuation boundaries and also split long newline blocks.
    chunks = re.split(r"(?<=[.!?])\s+|\n+", raw)
    out = []
    for chunk in chunks:
        c = _normalize_text(chunk).strip(" -•")
        if c:
            out.append(c)
    return out


def _sentence_relevance_score(sentence: str, topic_terms: list[str], domain: str = "") -> float:
    s = _normalize_text(sentence)
    low = s.lower()
    if not s:
        return 0.0

    score = 0.0
    if 80 <= len(s) <= 240:
        score += 0.18
    elif 55 <= len(s) <= 300:
        score += 0.1

    if topic_terms:
        match = sum(1 for t in topic_terms if t in low)
        ratio = match / max(len(topic_terms), 1)
        score += 0.45 * ratio
    else:
        score += 0.12

    if re.search(r"\b\d+(?:[.,]\d+)?\b", s):
        score += 0.08

    if any(k in low for k in ("araştırma", "çalışma", "study", "evidence", "kanıt", "risk", "benefit", "fayda", "hastalık", "sağlık", "beslenme", "davranış")):
        score += 0.1

    if domain.endswith((".gov", ".edu")):
        score += 0.08
    if any(h in domain for h in _TRUSTED_DOMAIN_HINTS):
        score += 0.07
    if any(h in domain for h in _LOW_VALUE_DOMAIN_HINTS):
        score -= 0.12

    return max(0.0, min(1.0, score))


def _extract_research_passages(text: str, topic_terms: list[str], domain: str = "", max_sentences: int = 8) -> list[str]:
    sentences = _split_sentences(text)
    if not sentences:
        return []

    ranked: list[tuple[float, str]] = []
    for sent in sentences:
        if _is_noise_sentence(sent):
            continue
        score = _sentence_relevance_score(sent, topic_terms, domain=domain)
        if score < 0.24:
            continue
        ranked.append((score, sent))

    if not ranked:
        return []

    ranked.sort(key=lambda x: x[0], reverse=True)
    picked: list[str] = []
    seen_keys: set[str] = set()
    for _, sent in ranked:
        key = _normalize_text(sent)[:70].lower()
        if key in seen_keys:
            continue
        seen_keys.add(key)
        picked.append(sent)
        if len(picked) >= max_sentences:
            break
    return picked


def _search_result_score(item: dict[str, Any], rank_index: int, total: int) -> float:
    url = str(item.get("url", "") or "").strip()
    title = _normalize_text(item.get("title", ""))
    snippet = _normalize_text(item.get("snippet", ""))
    domain = _domain_from_url(url)
    path = urlparse(url).path.lower() if url else ""

    score = max(0.0, 1.0 - (rank_index / max(total, 1)) * 0.35)

    if domain.endswith((".gov", ".edu")):
        score += 0.28
    if ".org" in domain:
        score += 0.07
    if any(h in domain for h in _TRUSTED_DOMAIN_HINTS):
        score += 0.22
    if any(h in domain for h in _LOW_VALUE_DOMAIN_HINTS):
        score -= 0.35
    if any(p in path for p in _LOW_VALUE_URL_PARTS):
        score -= 0.22

    if len(snippet) >= 90:
        score += 0.08
    elif len(snippet) < 25:
        score -= 0.08

    if not title:
        score -= 0.06

    return max(0.0, min(1.2, score))


def _quality_snapshot(sources: list["ResearchSource"], findings: list[str]) -> dict[str, Any]:
    total = len(sources)
    if total == 0:
        return {
            "total_sources": 0,
            "reliable_sources": 0,
            "reliability_pct": 0,
            "high_reliability": 0,
            "medium_reliability": 0,
            "low_reliability": 0,
            "finding_count": len(findings or []),
            "top_domains": [],
        }

    reliable = sum(1 for s in sources if float(s.reliability_score or 0) >= 0.6)
    high = sum(1 for s in sources if float(s.reliability_score or 0) >= 0.8)
    medium = sum(1 for s in sources if 0.6 <= float(s.reliability_score or 0) < 0.8)
    low = total - high - medium

    domains = [_domain_from_url(s.url) for s in sources if s.url]
    top_domains = [d for d, _ in Counter(domains).most_common(5) if d]

    return {
        "total_sources": total,
        "reliable_sources": reliable,
        "reliability_pct": int((reliable / total) * 100) if total else 0,
        "high_reliability": high,
        "medium_reliability": medium,
        "low_reliability": low,
        "finding_count": len(findings or []),
        "top_domains": top_domains,
    }


def _build_query_decomposition(topic: str) -> dict[str, Any]:
    clean = _normalize_text(topic)
    facets = ["definition", "history", "applications", "examples", "key claims"]
    base_terms = [token for token in _tokenize_topic(clean)[:6]]
    queries = [clean]
    for facet in facets:
        queries.append(f"{clean} {facet}")
    if base_terms:
        queries.extend([f"{clean} {' '.join(base_terms[:2])}", f"{clean} source pdf"])
    deduped = list(dict.fromkeys([item for item in queries if item]))
    return {
        "topic": clean,
        "facets": facets,
        "queries": deduped[:8],
        "topic_terms": base_terms,
    }


def _keyword_overlap_score(text: str, url: str, title: str, topic_terms: list[str]) -> float:
    hay = f"{_normalize_text(text)} {_normalize_text(title)} {url}".lower()
    if not hay:
        return 0.0
    hits = sum(1 for term in topic_terms if term in hay)
    return hits / max(len(topic_terms), 1)


def _build_research_contract_payload(topic: str, findings: list[str], sources: list["ResearchSource"]) -> dict[str, Any]:
    decomposition = _build_query_decomposition(topic)
    topic_terms = list(decomposition.get("topic_terms", []))
    claim_list: list[dict[str, Any]] = []
    citation_map: dict[str, list[dict[str, Any]]] = {}
    uncertainty_log: list[str] = []
    conflicts: list[dict[str, Any]] = []

    for idx, finding in enumerate(findings or [], start=1):
        claim_id = f"claim_{idx}"
        ranked_sources = sorted(
            sources or [],
            key=lambda source: (
                _keyword_overlap_score(finding, source.url, source.title, topic_terms),
                float(source.reliability_score or 0.0),
            ),
            reverse=True,
        )
        selected = []
        for source in ranked_sources[:3]:
            if not source.url:
                continue
            selected.append(
                {
                    "url": source.url,
                    "title": source.title,
                    "reliability_score": round(float(source.reliability_score or 0.0), 3),
                    "evidence_type": "summary",
                }
            )
        critical = idx <= 2
        source_urls = [item["url"] for item in selected]
        if critical and len(set(source_urls)) < 2:
            uncertainty_log.append(f"{claim_id}: kritik iddia icin ikinci bagimsiz kaynak eksik.")
        claim_list.append(
            {
                "claim_id": claim_id,
                "text": _compact_finding_text(finding, max_len=240),
                "source_urls": source_urls,
                "critical": critical,
                "confidence": round(
                    min(
                        1.0,
                        (sum(float(item["reliability_score"]) for item in selected) / max(len(selected), 1)) * (1.0 if not critical else 0.95),
                    ),
                    2,
                ),
            }
        )
        citation_map[claim_id] = selected

    numeric_claims: dict[str, set[str]] = {}
    for claim in claim_list:
        nums = set(re.findall(r"\b\d+(?:[.,]\d+)?\b", claim.get("text", "")))
        if nums:
            numeric_claims[claim["claim_id"]] = nums
    if len(numeric_claims) >= 2:
        all_nums = set()
        for nums in numeric_claims.values():
            all_nums.update(nums)
        if len(all_nums) > 1:
            conflicts.append(
                {
                    "type": "numeric_variation",
                    "claim_ids": list(numeric_claims.keys()),
                    "detail": "Sayisal ifadeler arasinda fark bulundu; manuel capraz kontrol onerilir.",
                }
            )

    return {
        "query_decomposition": decomposition,
        "claim_list": claim_list,
        "citation_map": citation_map,
        "critical_claim_ids": [claim["claim_id"] for claim in claim_list if claim.get("critical")],
        "conflicts": conflicts,
        "uncertainty_log": uncertainty_log,
    }


class ResearchDepth(Enum):
    QUICK = "quick"           # 2-3 kaynak, hızlı sonuç
    STANDARD = "standard"     # 5-7 kaynak, temel analiz
    COMPREHENSIVE = "comprehensive"  # 10-15 kaynak, detaylı analiz
    EXPERT = "expert"         # 15-20+ kaynak, akademik seviye


@dataclass
class ResearchSource:
    """Araştırma kaynağı"""
    url: str
    title: str
    snippet: str
    reliability_score: float = 0.0
    content: str = ""
    fetched: bool = False
    error: str | None = None

    def to_dict(self) -> dict:
        return {
            "url": self.url,
            "title": self.title,
            "snippet": self.snippet,
            "reliability_score": self.reliability_score,
            "fetched": self.fetched,
            "error": self.error
        }


@dataclass
class ResearchResult:
    """Araştırma sonucu"""
    id: str
    topic: str
    depth: ResearchDepth
    status: str = "pending"
    sources: list[ResearchSource] = field(default_factory=list)
    findings: list[str] = field(default_factory=list)
    summary: str = ""
    started_at: str | None = None
    completed_at: str | None = None
    progress: int = 0

    def to_dict(self) -> dict:
        return {
            "id": self.id,
            "topic": self.topic,
            "depth": self.depth.value,
            "status": self.status,
            "source_count": len(self.sources),
            "sources": [s.to_dict() for s in self.sources],
            "findings": self.findings,
            "summary": self.summary,
            "started_at": self.started_at,
            "completed_at": self.completed_at,
            "progress": self.progress
        }


# Store ongoing research
_research_tasks: dict[str, ResearchResult] = {}


def get_research_status(research_id: str) -> dict[str, Any]:
    """Get status of ongoing research"""
    if research_id not in _research_tasks:
        return {
            "success": False,
            "error": f"Araştırma bulunamadı: {research_id}"
        }

    result = _research_tasks[research_id]
    return {
        "success": True,
        "research_id": research_id,
        "status": result.status,
        "progress": result.progress,
        "topic": result.topic,
        "depth": result.depth.value,
        "source_count": len(result.sources),
        "finding_count": len(result.findings),
        "started_at": result.started_at,
        "completed_at": result.completed_at,
        "summary": result.summary if result.status == "completed" else None
    }


def get_research_result_snapshot(research_id: str) -> dict[str, Any]:
    """Get research snapshot (completed-only strict mode)."""
    if research_id not in _research_tasks:
        return {
            "success": False,
            "error": f"Araştırma bulunamadı: {research_id}"
        }

    result = _research_tasks[research_id]

    if result.status != "completed":
        return {
            "success": False,
            "error": f"Araştırma henüz tamamlanmadı. Durum: {result.status}",
            "progress": result.progress
        }

    sources_dict = [s.to_dict() for s in result.sources]
    return {
        "success": True,
        "research_id": research_id,
        "topic": result.topic,
        "depth": result.depth.value,
        "sources": sources_dict,
        "findings": result.findings,
        "summary": result.summary,
        "completed_at": result.completed_at
    }


async def advanced_research(
    topic: str,
    depth: str = "standard",
    sources: list[str] | None = None,
    language: str = "tr",
    include_evaluation: bool = True,
    generate_report: bool = False,  # CHANGED: Report generation disabled for speed (task takes 5 mins otherwise)
    source_policy: str = "balanced",
    min_reliability: float = 0.55,
    max_findings: int = 6,
    citation_style: str = "none",
    include_bibliography: bool = False,
) -> dict[str, Any]:
    """
    Gelişmiş araştırma yap ve rapor oluştur

    Args:
        topic: Araştırma konusu
        depth: Derinlik seviyesi ("quick", "standard", "comprehensive", "expert")
        sources: Belirli kaynaklar (URL listesi)
        language: Arama dili ("tr", "en")
        include_evaluation: Kaynak güvenilirlik değerlendirmesi dahil et
        generate_report: Profesyonel rapor oluştur
        source_policy: Kaynak seçim politikası ("balanced", "trusted", "academic", "official")
        min_reliability: Özet ve bulgular için minimum güvenilirlik eşiği (0-1)
        max_findings: Döndürülecek maksimum bulgu sayısı

    Returns:
        dict: Araştırma sonuçları ve rapor yolu
    """
    try:
        if not topic or not topic.strip():
            return {"success": False, "error": "Araştırma konusu gerekli"}

        topic = topic.strip()
        source_policy = _normalize_source_policy(source_policy)
        citation_style = str(citation_style or "none").strip().lower()
        if citation_style not in {"none", "apa7", "mla", "ieee", "chicago"}:
            citation_style = "none"
        include_bibliography = bool(include_bibliography)
        try:
            min_reliability = max(0.0, min(1.0, float(min_reliability)))
        except Exception:
            min_reliability = 0.55
        try:
            max_findings = max(3, min(12, int(max_findings)))
        except Exception:
            max_findings = 6

        # Parse depth
        try:
            research_depth = ResearchDepth(depth.lower())
        except ValueError:
            research_depth = ResearchDepth.STANDARD

        # Determine source count based on depth
        source_counts = {
            ResearchDepth.QUICK: 3,
            ResearchDepth.STANDARD: 6,
            ResearchDepth.COMPREHENSIVE: 12,
            ResearchDepth.EXPERT: 18
        }
        target_sources = source_counts[research_depth]
        eval_cap_map = {
            ResearchDepth.QUICK: 2,
            ResearchDepth.STANDARD: 6,
            ResearchDepth.COMPREHENSIVE: 10,
            ResearchDepth.EXPERT: 14,
        }
        eval_concurrency_map = {
            ResearchDepth.QUICK: 2,
            ResearchDepth.STANDARD: 4,
            ResearchDepth.COMPREHENSIVE: 5,
            ResearchDepth.EXPERT: 6,
        }
        eval_timeout_map = {
            ResearchDepth.QUICK: 8,
            ResearchDepth.STANDARD: 10,
            ResearchDepth.COMPREHENSIVE: 12,
            ResearchDepth.EXPERT: 15,
        }

        # Create research ID
        research_id = f"research_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:6]}"

        result = ResearchResult(
            id=research_id,
            topic=topic,
            depth=research_depth,
            status="running",
            started_at=datetime.now().isoformat()
        )
        _research_tasks[research_id] = result

        logger.info(f"Araştırma başlatıldı: {topic} ({research_depth.value})")

        try:
            # Step 1: Web search for sources
            result.progress = 10
            # Pull larger pool and then filter by source policy.
            search_pool_size = max(target_sources * 2, target_sources + 4)
            search_results = await _perform_web_search(topic, search_pool_size, language)
            search_results = _apply_source_policy(search_results, source_policy, target_sources=target_sources)
            policy_stats = _source_policy_stats(search_results, source_policy)

            if not search_results:
                result.status = "failed"
                result.completed_at = datetime.now().isoformat()
                return {
                    "success": False,
                    "research_id": research_id,
                    "error": "Arama sonucu bulunamadı"
                }

            # Add custom sources if provided
            if sources:
                for url in sources[:5]:  # Limit custom sources
                    search_results.append({
                        "url": url,
                        "title": "Kullanıcı kaynağı",
                        "snippet": ""
                    })

            # Create source objects
            for sr in search_results[:target_sources]:
                source = ResearchSource(
                    url=sr.get("url", ""),
                    title=sr.get("title", ""),
                    snippet=sr.get("snippet", ""),
                    reliability_score=float(sr.get("_rank_score", 0.5) or 0.5),
                )
                result.sources.append(source)

            result.progress = 30

            # Step 2: Fetch and evaluate sources (CONCURRENT for speed)
            if include_evaluation:
                timeout_s = int(eval_timeout_map.get(research_depth, 10))
                eval_cap = int(eval_cap_map.get(research_depth, target_sources))
                sources_to_eval = result.sources[:eval_cap]

                async def _evaluate_with_timeout(source):
                    try:
                        # Per-source timeout is adaptive by depth.
                        eval_result = await asyncio.wait_for(
                            evaluate_source(source.url),
                            timeout=timeout_s
                        )
                        if eval_result.get("success"):
                            source.reliability_score = eval_result.get("reliability_score", 0.5)
                            source.content = eval_result.get("content_preview", "")
                            source.fetched = True
                        else:
                            source.error = eval_result.get("error")
                    except asyncio.TimeoutError:
                        source.error = "Timeout"
                    except Exception as e:
                        source.error = str(e)

                # Run evaluations concurrently with adaptive limits per depth.
                semaphore = asyncio.Semaphore(int(eval_concurrency_map.get(research_depth, 3)))
                async def _with_semaphore(source):
                    async with semaphore:
                        await _evaluate_with_timeout(source)

                await asyncio.gather(*[_with_semaphore(s) for s in sources_to_eval], return_exceptions=True)
                result.progress = 70

            # Filter weakest sources after scoring, but keep enough context.
            result.sources = _apply_min_reliability(result.sources, min_reliability=min_reliability, keep_at_least=2)

            # Step 3: Extract findings
            result.progress = 75
            result.findings = await _extract_findings(result.sources, topic, max_findings=max_findings)

            # Step 4: Generate summary
            result.progress = 90
            result.summary = await _generate_summary(
                topic,
                result.findings,
                result.sources,
                source_policy=source_policy,
                min_reliability=min_reliability,
                source_policy_stats=policy_stats,
            )
            quality = _quality_snapshot(result.sources, result.findings)
            research_contract = _build_research_contract_payload(topic, result.findings, result.sources)

            # Step 5: Generate professional report with visualizations
            result.progress = 95
            sources_dict = [s.to_dict() for s in result.sources]

            report_data = {
                "sources": sources_dict,
                "findings": result.findings,
                "summary": result.summary,
                "research_contract": research_contract,
                "depth": research_depth.value,
                "source_policy": source_policy,
                "min_reliability": min_reliability,
                "source_policy_stats": policy_stats,
            }

            report_paths = []
            quick_report_path = _persist_quick_research_report(topic, result)
            if quick_report_path:
                report_paths.append(quick_report_path)
            if generate_report:
                try:
                    from .advanced_report import generate_advanced_professional_report

                    # Generate advanced professional report with visualizations
                    report_result = await generate_advanced_professional_report(topic, report_data)
                    if report_result.get("success"):
                        report_paths.append(report_result.get("path"))
                        logger.info(f"Profesyonel rapor oluşturuldu: {report_result.get('path')}")
                        logger.info(f"Rapor metrikleri - Kapsam: {report_result.get('metrics', {}).get('coverage'):.0f}%, "
                                  f"Güvenilirlik: {report_result.get('metrics', {}).get('reliability'):.0f}%, "
                                  f"Bütünlük: {report_result.get('metrics', {}).get('completeness'):.0f}%")
                    else:
                        logger.warning(f"Rapor oluşturma başarısız: {report_result.get('error')}")

                except Exception as e:
                    logger.warning(f"Gelişmiş rapor oluşturulamadı, temel format kullanılıyor: {e}")
                    try:
                        from .report_generator import create_research_report
                        pdf_result = await create_research_report(topic, report_data, format="pdf")
                        if pdf_result.get("success"):
                            report_paths.append(pdf_result.get("path"))
                            logger.info(f"Temel PDF rapor oluşturuldu: {pdf_result.get('path')}")
                    except Exception as fallback_error:
                        logger.warning(f"Fallback rapor oluşturulamadı: {fallback_error}")

            result.status = "completed"
            result.completed_at = datetime.now().isoformat()
            result.progress = 100

            logger.info(f"Araştırma tamamlandı: {topic} - {len(result.sources)} kaynak, {len(report_paths)} rapor")

            # Format professional report message for Telegram
            message = f"📊 **{topic.upper()} - ARAŞTIRMA RAPORU**\n"
            message += f"{'='*50}\n\n"
            message += f"**📈 Özet:**\n{result.summary}\n\n"
            message += f"**🔍 Ana Bulgular ({len(result.findings)}):**\n"
            for i, finding in enumerate(result.findings, 1):
                message += f"{i}. {finding}\n"
            message += f"\n**📚 Kaynaklar ({len(result.sources)}):**\n"
            for i, src in enumerate(result.sources[:8], 1):  # Limit to 8 sources for Telegram
                reliability = f"({src.reliability_score*100:.0f}% güvenilir)" if src.reliability_score > 0 else ""
                message += f"{i}. [{src.title}]({src.url}) {reliability}\n"

            if report_paths:
                message += f"\n**📄 Detaylı Rapor:**\n"
                for path in report_paths:
                    message += f"- {path}\n"

            # Offer detailed report generation
            message += f"\n{'='*50}\n"
            message += f"💾 **Detaylı profesyonel rapor (DOCX) oluşturmak ister misin?**\n"
            message += f"Komut: `Raporla` veya `araştırma raporu oluştur`\n"

            # Store last research topic in global cache for report generation
            try:
                global _last_research_topic
                _last_research_topic = topic
                global _last_research_result
                _last_research_result = {
                    "research_id": research_id,
                    "topic": topic,
                    "findings": list(result.findings),
                    "summary": result.summary,
                    "sources": sources_dict,
                    "quality": quality,
                    "research_contract": research_contract,
                    "source_policy": source_policy,
                    "min_reliability": min_reliability,
                    "source_policy_stats": policy_stats,
                    "citation_style": citation_style,
                    "include_bibliography": include_bibliography,
                    "report_paths": list(report_paths),
                    "completed_at": result.completed_at,
                }
                logger.info(f"Stored last research topic in cache: {topic}")
            except Exception as e:
                logger.debug(f"Could not store topic in cache: {e}")

            return {
                "success": True,
                "research_id": research_id,
                "topic": topic,
                "depth": research_depth.value,
                "source_policy": source_policy,
                "min_reliability": min_reliability,
                "source_policy_stats": policy_stats,
                "citation_style": citation_style,
                "include_bibliography": include_bibliography,
                "source_count": len(result.sources),
                "sources": sources_dict,
                "findings": result.findings,
                "summary": result.summary,
                "quality": quality,
                "research_contract": research_contract,
                "report_paths": report_paths,
                "message": message
            }

        except Exception as e:
            result.status = "failed"
            result.completed_at = datetime.now().isoformat()
            logger.error(f"Araştırma hatası: {e}")
            return {
                "success": False,
                "research_id": research_id,
                "error": f"Araştırma başarısız: {str(e)}"
            }

    except Exception as e:
        logger.error(f"Gelişmiş araştırma hatası: {e}")
        return {"success": False, "error": f"Araştırma yapılamadı: {str(e)}"}


async def evaluate_source(
    url: str,
    criteria: dict[str, bool] | None = None
) -> dict[str, Any]:
    """
    Kaynak güvenilirlik değerlendirmesi

    Args:
        url: Değerlendirilecek URL
        criteria: Değerlendirme kriterleri

    Returns:
        dict: Değerlendirme sonucu
    """
    try:
        if not url:
            return {"success": False, "error": "URL gerekli"}

        # Default criteria
        if criteria is None:
            criteria = {
                "check_domain": True,
                "check_https": True,
                "check_content": True
            }

        score = 0.5  # Base score
        factors = []

        # Domain reliability scoring
        parsed = urlparse(url)
        domain = parsed.netloc.lower()

        # HTTPS check
        if criteria.get("check_https", True):
            if parsed.scheme == "https":
                score += 0.1
                factors.append({"name": "HTTPS", "value": True, "weight": 0.1})
            else:
                score -= 0.1
                factors.append({"name": "HTTPS", "value": False, "weight": -0.1})

        # Domain reliability
        if criteria.get("check_domain", True):
            # Academic and government domains
            trusted_tlds = [".edu", ".gov", ".ac."]
            academic_domains = ["arxiv.org", "scholar.google", "researchgate.net", "academia.edu"]
            news_reliable = ["bbc.", "reuters.", "apnews.", "npr.org"]
            unreliable_patterns = ["blogspot.", "wordpress.", "tumblr.", "medium.com", "pinterest.", "reddit.com"]

            for tld in trusted_tlds:
                if tld in domain:
                    score += 0.15
                    factors.append({"name": f"TLD {tld}", "value": True, "weight": 0.15})
                    break

            for ad in academic_domains:
                if ad in domain:
                    score += 0.2
                    factors.append({"name": "Academic domain", "value": True, "weight": 0.2})
                    break

            for nr in news_reliable:
                if nr in domain:
                    score += 0.1
                    factors.append({"name": "Reliable news", "value": True, "weight": 0.1})
                    break

            for up in unreliable_patterns:
                if up in domain:
                    score -= 0.15
                    factors.append({"name": "User-generated content", "value": True, "weight": -0.15})
                    break

            if any(td in domain for td in _TRUSTED_DOMAIN_HINTS):
                score += 0.12
                factors.append({"name": "Trusted source hint", "value": True, "weight": 0.12})
            if any(ld in domain for ld in _LOW_VALUE_DOMAIN_HINTS):
                score -= 0.15
                factors.append({"name": "Low-value source hint", "value": True, "weight": -0.15})

        # Try to fetch content preview
        content_preview = ""
        if criteria.get("check_content", True):
            try:
                from tools.web_tools import fetch_page
                fetch_result = await fetch_page(url, extract_content=True)
                if fetch_result.get("success"):
                    content = _normalize_text(fetch_result.get("content", ""))
                    title_terms = _tokenize_topic(str(fetch_result.get("title", "")))
                    passages = _extract_research_passages(content, title_terms, domain=domain, max_sentences=10)
                    if passages:
                        content_preview = " ".join(passages)[:2000]
                    elif content:
                        content_preview = content[:1200]

                    # Content quality indicators
                    if len(content) > 1000:
                        score += 0.05
                        factors.append({"name": "Content length", "value": "long", "weight": 0.05})
                    if len(content) > 2500:
                        score += 0.05
                        factors.append({"name": "Content depth", "value": "detailed", "weight": 0.05})
                    if passages and len(passages) >= 3:
                        score += 0.06
                        factors.append({"name": "Extractable passages", "value": len(passages), "weight": 0.06})
                    if re.search(r"\b\d+(?:[.,]\d+)?\b", content):
                        score += 0.03
                        factors.append({"name": "Has quantitative signals", "value": True, "weight": 0.03})

                    # Check for citations/references
                    citation_keywords = ["source:", "reference", "citation", "et al.", "study", "research"]
                    if any(kw in content.lower() for kw in citation_keywords):
                        score += 0.1
                        factors.append({"name": "Has citations", "value": True, "weight": 0.1})
                    noise_hits = sum(1 for p in _NOISE_PATTERNS if p in content.lower())
                    if noise_hits >= 4:
                        score -= 0.07
                        factors.append({"name": "High noise content", "value": noise_hits, "weight": -0.07})
            except Exception as e:
                factors.append({"name": "Content fetch", "value": False, "error": str(e)})

        # Normalize score
        reliability_score = max(0.0, min(1.0, score))

        # Determine reliability level
        if reliability_score >= 0.8:
            level = "high"
        elif reliability_score >= 0.6:
            level = "medium"
        elif reliability_score >= 0.4:
            level = "low"
        else:
            level = "very_low"

        return {
            "success": True,
            "url": url,
            "domain": domain,
            "reliability_score": round(reliability_score, 2),
            "reliability_level": level,
            "factors": factors,
            "content_preview": content_preview
        }

    except Exception as e:
        logger.error(f"Kaynak değerlendirme hatası: {e}")
        return {"success": False, "error": f"Kaynak değerlendirilemedi: {str(e)}"}


async def quick_research(
    topic: str,
    max_sources: int = 3
) -> dict[str, Any]:
    """
    Hızlı araştırma - minimal kaynakla hızlı sonuç

    Args:
        topic: Araştırma konusu
        max_sources: Maksimum kaynak sayısı

    Returns:
        dict: Hızlı araştırma sonuçları
    """
    return await advanced_research(
        topic=topic,
        depth="quick",
        include_evaluation=False
    )


async def _perform_web_search(
    query: str,
    num_results: int,
    language: str
) -> list[dict]:
    """Web araması yap"""
    try:
        from tools.web_tools import web_search
        target_n = max(int(num_results or 1), 1)
        result = await web_search(query, num_results=max(target_n * 2, 8), language=language)
        if result.get("success"):
            raw = result.get("results", []) or []
            deduped: list[dict[str, Any]] = []
            seen_urls: set[str] = set()
            seen_titles: set[str] = set()
            for item in raw:
                if not isinstance(item, dict):
                    continue
                url = str(item.get("url", "") or "").strip()
                if not url:
                    continue
                title = str(item.get("title", "") or "").strip().lower()
                url_key = url.split("#", 1)[0].rstrip("/")
                title_key = title[:120]
                if url_key in seen_urls:
                    continue
                if title_key and title_key in seen_titles:
                    continue
                seen_urls.add(url_key)
                if title_key:
                    seen_titles.add(title_key)
                deduped.append(item)

            ranked: list[tuple[float, dict[str, Any]]] = []
            total = len(deduped)
            for idx, item in enumerate(deduped):
                score = _search_result_score(item, idx, total)
                ranked.append((score, item))

            ranked.sort(key=lambda x: x[0], reverse=True)
            final_results: list[dict[str, Any]] = []
            for score, item in ranked:
                domain = _domain_from_url(str(item.get("url", "") or ""))
                if score < 0.20 and len(final_results) >= min(3, target_n):
                    continue
                if any(h in domain for h in _LOW_VALUE_DOMAIN_HINTS) and len(final_results) >= min(3, target_n):
                    continue
                item_with_score = dict(item)
                item_with_score["_rank_score"] = round(float(score), 3)
                final_results.append(item_with_score)
                if len(final_results) >= target_n:
                    break

            if not final_results:
                final_results = []
                for score, item in ranked[:target_n]:
                    item_with_score = dict(item)
                    item_with_score["_rank_score"] = round(float(score), 3)
                    final_results.append(item_with_score)
            return final_results
        return []
    except Exception as e:
        logger.warning(f"Web arama hatası: {e}")
        return []


async def _extract_findings(
    sources: list[ResearchSource],
    topic: str,
    max_findings: int = 10,
) -> list[str]:
    """Kaynaklardan yüksek kaliteli, dedupe edilmiş ve konuya alakalı bulgular çıkar."""
    topic_terms = _tokenize_topic(topic)
    findings: list[str] = []
    seen_keys: set[str] = set()
    candidates: list[tuple[float, str, str, float]] = []

    for source in sources:
        domain = _domain_from_url(source.url)
        source_rel = max(0.0, min(1.0, float(source.reliability_score or 0.0)))
        body = _normalize_text(source.content if source.fetched and source.content else source.snippet)
        if not body:
            continue

        passages = _extract_research_passages(
            body,
            topic_terms=topic_terms,
            domain=domain,
            max_sentences=3,
        )

        if not passages and source.snippet:
            snippet = _normalize_text(source.snippet)
            if not _is_noise_sentence(snippet):
                passages = [snippet]

        for passage in passages:
            clean = _compact_finding_text(_normalize_text(passage).lstrip("•- "))
            if not clean:
                continue
            if len(clean) < 35:
                continue
            if not re.search(r"[a-zA-ZğüşöçıİĞÜŞÖÇ]{8,}", clean):
                continue
            key = clean[:90].lower()
            if key in seen_keys:
                continue
            seen_keys.add(key)

            relevance = _sentence_relevance_score(clean, topic_terms, domain=domain)
            composite = (relevance * 0.75) + (source_rel * 0.25)
            if any(p in clean.lower() for p in _NOISE_PATTERNS):
                composite -= 0.2
            candidates.append((max(0.0, composite), clean, domain, source_rel))

    if candidates:
        candidates.sort(key=lambda x: x[0], reverse=True)
        domain_counts: dict[str, int] = {}
        findings.clear()
        for score, clean, domain, source_rel in candidates:
            _ = score
            if domain and domain_counts.get(domain, 0) >= 2:
                continue
            confidence_pct = int(round(max(0.0, min(1.0, source_rel)) * 100))
            if domain:
                domain_counts[domain] = domain_counts.get(domain, 0) + 1
                findings.append(f"• {clean} (Kaynak: {domain}, Güven: %{confidence_pct})")
            else:
                findings.append(f"• {clean} (Güven: %{confidence_pct})")
            if len(findings) >= max(3, int(max_findings or 10)):
                return findings

    # Ultra fallback: en azından snippet tabanlı 3 bulgu üret.
    if len(findings) < 3:
        seen_keys = {str(item).lower()[:90] for item in findings}
        for source in sources:
            snippet = _normalize_text(source.snippet)
            if not snippet:
                continue
            if _is_noise_sentence(snippet):
                continue
            snippet = _compact_finding_text(snippet)
            key = snippet[:90].lower()
            if key in seen_keys:
                continue
            seen_keys.add(key)
            domain = _domain_from_url(source.url)
            findings.append(f"• {snippet}" + (f" (Kaynak: {domain})" if domain else ""))
            if len(findings) >= 5:
                break

    return findings[: max(3, int(max_findings or 10))]


def _suggest_followup_actions(topic: str, findings: list[str], reliability_pct: int) -> list[str]:
    low_topic = _normalize_text(topic).lower()
    has_health = any(k in low_topic for k in ("sağlık", "saglik", "hastalık", "hastalik", "beslenme", "veteriner"))
    has_business = any(k in low_topic for k in ("e-ticaret", "eticaret", "satış", "satis", "stok", "sipariş", "siparis"))
    has_tech = any(k in low_topic for k in ("yazılım", "yazilim", "kod", "api", "framework", "model"))

    actions: list[str] = []
    if reliability_pct < 70:
        actions.append("Kaynak havuzunu akademik/resmi filtre ile yeniden çalıştır.")
    else:
        actions.append("En yüksek güvenilir 3 kaynağı referans alarak kısa karar notu çıkar.")

    if has_health:
        actions.append("Bulguları klinik/uzman doğrulaması gerektiren başlıklar için ayrı etiketle.")
    elif has_business:
        actions.append("Bulgulardan KPI etkisi olan maddeleri tabloya dönüştür (hacim, risk, fırsat).")
    elif has_tech:
        actions.append("Teknik bulgular için uygulanabilir backlog maddeleri (P0/P1/P2) üret.")
    else:
        actions.append("Bulguları tema bazında gruplayıp eylem planına dönüştür.")

    if findings:
        actions.append("Çelişkili görünen maddeleri kaynak karşılaştırması ile netleştir.")
    return actions[:3]


async def _generate_summary(
    topic: str,
    findings: list[str],
    sources: list[ResearchSource],
    source_policy: str = "balanced",
    min_reliability: float = 0.55,
    source_policy_stats: dict[str, Any] | None = None,
) -> str:
    """Araştırma sonuçlarını profesyonel, kanıt odaklı bir özet metnine dönüştür."""
    quality = _quality_snapshot(sources, findings)
    source_count = quality["total_sources"]
    reliable_count = quality["reliable_sources"]
    reliability_pct = quality["reliability_pct"]
    high_rel = quality["high_reliability"]
    med_rel = quality["medium_reliability"]
    low_rel = quality["low_reliability"]
    top_domains = quality["top_domains"]

    cleaned_findings: list[str] = []
    seen: set[str] = set()
    for finding in findings:
        item = _normalize_text(str(finding or "").lstrip("•- "))
        if len(item) < 24:
            continue
        key = item[:90].lower()
        if key in seen:
            continue
        seen.add(key)
        cleaned_findings.append(item)

    avg_reliability = 0.0
    if sources:
        avg_reliability = sum(float(s.reliability_score or 0.0) for s in sources) / max(len(sources), 1)
    confidence_label = "yüksek" if reliability_pct >= 80 else ("orta" if reliability_pct >= 60 else "düşük")
    actions = _suggest_followup_actions(topic, cleaned_findings, reliability_pct)

    lines: list[str] = []
    lines.append("Yönetici Özeti:")
    lines.append(
        f"'{topic}' araştırmasında {source_count} kaynak incelendi; güvenilirlik eşiğini geçen kaynak oranı %{reliability_pct} ({confidence_label} güven)."
    )
    lines.append("")
    lines.append(f"Araştırma Konusu: {topic}")
    lines.append(f"Kapsam: {source_count} kaynak tarandı; güvenilirlik eşiğini geçen kaynak {reliable_count}/{source_count} (%{reliability_pct}).")
    lines.append(
        f"Kaynak Profili: yüksek={high_rel}, orta={med_rel}, düşük={low_rel}, ortalama güven={avg_reliability:.2f}."
    )
    if source_policy != "balanced":
        lines.append(f"Kaynak Politikası: {source_policy} (min güvenilirlik: {min_reliability:.2f}).")
        stats = source_policy_stats or {}
        if bool(stats.get("fallback_used")):
            lines.append("Not: Seçilen politika için yeterli kaynak bulunamadığından sınırlı fallback uygulandı.")
    if top_domains:
        lines.append("Öne Çıkan Alan Adları: " + ", ".join(top_domains[:5]))

    if cleaned_findings:
        lines.append("")
        lines.append("Ana Bulgular (kanıt odaklı):")
        for idx, finding in enumerate(cleaned_findings[:6], 1):
            lines.append(f"{idx}. {finding}")
    else:
        lines.append("")
        lines.append("Ana Bulgular: Bu taramada doğrudan kullanılabilir bulgu sınırlı kaldı.")

    lines.append("")
    lines.append("Kanıt Matrisi (ilk 5 kaynak):")
    ranked_sources = sorted(sources, key=lambda s: float(s.reliability_score or 0.0), reverse=True)[:5]
    if ranked_sources:
        for idx, src in enumerate(ranked_sources, 1):
            domain = _domain_from_url(src.url)
            title = _normalize_text(src.title) or domain or "Başlıksız kaynak"
            rel = float(src.reliability_score or 0.0)
            lines.append(f"{idx}. {title} | {domain} | güven={rel:.2f}")
    else:
        lines.append("- Kaynak bilgisi sınırlı.")

    lines.append("")
    lines.append("Operasyonel Öneriler:")
    if actions:
        for item in actions:
            lines.append(f"- {item}")
    else:
        lines.append("- Ek öneri üretilemedi.")

    lines.append("")
    lines.append("Sınırlılıklar:")
    lines.append("- Bulgular web kaynaklarından otomatik çıkarılmıştır; alan uzmanı doğrulaması önerilir.")
    lines.append("- Tarih/bölge bilgisi olmayan kaynaklarda genelleme yapılmamalıdır.")
    lines.append("")
    lines.append("Önerilen Devam Adımı: İstersen bu başlığı alt konulara bölüp (sağlık, beslenme, eğitim, ırklar) karşılaştırmalı rapor üretebilirim.")

    return "\n".join(lines)


def get_research_result(research_id: str) -> dict[str, Any]:
    """Araştırma sonucunu getir (yalnızca tamamlandıysa)."""
    if research_id not in _research_tasks:
        return {"success": False, "error": f"Araştırma bulunamadı: {research_id}"}

    result = _research_tasks[research_id]
    if result.status != "completed":
        return {
            "success": False,
            "error": f"Araştırma henüz tamamlanmadı. Durum: {result.status}",
            "progress": result.progress,
        }

    return {
        "success": True,
        **result.to_dict()
    }


def get_last_research_result() -> dict[str, Any]:
    """Son tamamlanan araştırma özetini döndürür."""
    if not _last_research_result:
        return {"success": False, "error": "Önce tamamlanmış bir araştırma gerekli."}
    return {"success": True, "data": dict(_last_research_result)}


def _persist_quick_research_report(topic: str, result: ResearchResult) -> str:
    """Always persist a lightweight markdown report for reliability and handoff."""
    try:
        day = datetime.now().strftime("%Y%m%d")
        ts = datetime.now().strftime("%H%M%S")
        safe_topic = "".join(c if c.isalnum() or c in " -_" else "_" for c in topic).strip()[:64] or "research"
        out_dir = RESEARCH_REPORT_DIR / day
        out_dir.mkdir(parents=True, exist_ok=True)
        out_path = out_dir / f"{safe_topic}_{ts}.md"

        lines = [
            f"# Araştırma: {topic}",
            "",
            f"- Derinlik: {result.depth.value}",
            f"- Kaynak sayısı: {len(result.sources)}",
            f"- Tamamlanma: {result.completed_at or datetime.now().isoformat()}",
            "",
            "## Özet",
            result.summary or "Özet üretilemedi.",
            "",
            "## Bulgular",
        ]

        if result.findings:
            for finding in result.findings[:12]:
                clean = str(finding or "").lstrip("•- ").strip()
                if clean:
                    lines.append(f"- {clean}")
        else:
            lines.append("- Bulgu bulunamadı.")

        lines.extend(["", "## Kaynaklar"])
        for idx, source in enumerate(result.sources[:15], start=1):
            title = str(source.title or "Başlıksız kaynak").strip()
            url = str(source.url or "").strip()
            rel = f"{float(source.reliability_score or 0.0):.2f}"
            lines.append(f"{idx}. {title} ({rel})")
            if url:
                lines.append(f"   - {url}")

        out_path.write_text("\n".join(lines), encoding="utf-8")
        return str(out_path)
    except Exception as e:
        logger.warning(f"Hızlı araştırma raporu yazılamadı: {e}")
        return ""


async def save_research_to_document(
    topic: str,
    findings: list[str],
    sources: list[dict],
    output_path: str = None,
    output_format: str = "docx"
) -> dict[str, Any]:
    """
    Araştırma sonuçlarını belge olarak kaydet
    
    Args:
        topic: Araştırma konusu
        findings: Bulgular listesi
        sources: Kaynak listesi
        output_path: Çıktı dosyası yolu (opsiyonel)
        output_format: Çıktı formatı (docx, txt, md)
    
    Returns:
        dict: Kayıt sonucu
    """
    from datetime import datetime
    from pathlib import Path
    
    try:
        # Default output path
        if not output_path:
            from config.settings import HOME_DIR
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            safe_topic = "".join(c for c in topic if c.isalnum() or c in " -_")[:50]
            filename = f"arastirma_{safe_topic}_{timestamp}.{output_format}"
            output_path = str(HOME_DIR / "Desktop" / filename)
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        if output_format == "docx":
            # Word belgesi oluştur
            try:
                from docx import Document
                from docx.shared import Pt, Inches
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                
                doc = Document()
                
                # Başlık
                title = doc.add_heading(f"Araştırma Raporu: {topic}", level=0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Tarih
                doc.add_paragraph(f"Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
                doc.add_paragraph()
                
                # Özet
                doc.add_heading("Özet", level=1)
                doc.add_paragraph(f"Bu rapor '{topic}' konusunda {len(sources)} kaynak incelenerek hazırlanmıştır.")
                
                # Bulgular
                doc.add_heading("Önemli Bulgular", level=1)
                for finding in findings:
                    # Bullet point olarak ekle
                    p = doc.add_paragraph(style='List Bullet')
                    finding_clean = finding.lstrip("• -").strip()
                    p.add_run(finding_clean)
                
                # Kaynaklar
                doc.add_heading("Kaynaklar", level=1)
                for i, source in enumerate(sources, 1):
                    title_text = source.get("title", "Başlıksız")
                    url = source.get("url", "")
                    p = doc.add_paragraph()
                    p.add_run(f"{i}. {title_text}").bold = True
                    if url:
                        p.add_run(f"\n   {url}")
                
                doc.save(str(output_path))
                
            except ImportError:
                # python-docx yoksa txt olarak kaydet
                output_format = "txt"
                output_path = output_path.with_suffix(".txt")
        
        if output_format in ["txt", "md"]:
            # Metin dosyası olarak kaydet
            content_lines = [
                f"# Araştırma Raporu: {topic}",
                f"Tarih: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
                "",
                "## Özet",
                f"Bu rapor '{topic}' konusunda {len(sources)} kaynak incelenerek hazırlanmıştır.",
                "",
                "## Önemli Bulgular",
            ]
            
            for finding in findings:
                finding_clean = finding.lstrip("• -").strip()
                content_lines.append(f"- {finding_clean}")
            
            content_lines.extend([
                "",
                "## Kaynaklar",
            ])
            
            for i, source in enumerate(sources, 1):
                title_text = source.get("title", "Başlıksız")
                url = source.get("url", "")
                content_lines.append(f"{i}. **{title_text}**")
                if url:
                    content_lines.append(f"   {url}")
            
            output_path.write_text("\n".join(content_lines), encoding="utf-8")
        
        logger.info(f"Araştırma belgesi kaydedildi: {output_path}")
        
        return {
            "success": True,
            "path": str(output_path),
            "format": output_format,
            "message": f"Araştırma raporu kaydedildi: {output_path.name}"
        }
        
    except Exception as e:
        logger.error(f"Belge kaydetme hatası: {e}")
        return {"success": False, "error": f"Belge kaydedilemedi: {str(e)}"}
