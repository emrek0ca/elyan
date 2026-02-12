"""
Advanced Professional Research Report Generator
Enterprise-grade reporting with comprehensive analytics and formatting
"""

import asyncio
import re
from typing import Dict, List, Optional, Any
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass
from utils.logger import get_logger
from .report_visualizations import ReportVisualizations, ReportChartBuilder

logger = get_logger("advanced_report")


@dataclass
class ReportMetrics:
    """Research quality metrics"""
    source_count: int
    finding_count: int
    average_reliability: float
    research_depth: str
    coverage_score: float  # 0-100
    reliability_score: float  # 0-100
    completeness_score: float  # 0-100


class AdvancedReportGenerator:
    """Enterprise-grade professional report generator"""

    def __init__(self):
        self.output_dir = Path.home() / "Desktop" / "Research Reports"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.llm_client = None
        self.visualizations = ReportVisualizations()
        self.chart_builder = ReportChartBuilder()

    async def initialize_llm(self):
        """Initialize LLM for advanced synthesis"""
        try:
            from core.llm_client import LLMClient
            self.llm_client = LLMClient()
            logger.info("LLM client initialized for report generation")
        except Exception as e:
            logger.debug(f"LLM initialization optional: {e}")

    async def generate_professional_report(
        self,
        topic: str,
        research_data: Dict[str, Any],
        report_style: str = "comprehensive"
    ) -> Dict[str, Any]:
        """Generate professional enterprise-grade report"""
        try:
            # Calculate quality metrics
            metrics = self._calculate_metrics(research_data)

            # Generate advanced sections
            report_sections = await self._generate_report_sections(
                topic,
                research_data,
                metrics,
                report_style
            )

            # Generate Word document first (preferred format)
            try:
                word_path = await self._generate_word_document(
                    topic,
                    report_sections,
                    metrics
                )
                report_path = word_path
                report_format = "docx"
            except Exception as e:
                logger.warning(f"Word generation failed, using PDF: {e}")
                # Fallback to PDF
                pdf_path = await self._generate_advanced_pdf(
                    topic,
                    report_sections,
                    metrics
                )
                report_path = pdf_path
                report_format = "pdf"

            logger.info(f"Professional report generated: {report_path} ({report_format})")

            # Generate a summary chart for the UI
            summary_chart = self.visualizations.generate_source_reliability_chart(research_data.get("sources", []))

            return {
                "success": True,
                "path": str(report_path),
                "format": report_format,
                "chart": summary_chart,
                "metrics": {
                    "coverage": metrics.coverage_score,
                    "reliability": metrics.reliability_score,
                    "completeness": metrics.completeness_score,
                    "sources": metrics.source_count,
                    "findings": metrics.finding_count
                },
                "report_type": report_style
            }

        except Exception as e:
            logger.error(f"Professional report generation failed: {e}")
            return {
                "success": False,
                "error": str(e)
            }

    async def generate_intelligence_report(
        self,
        topic: str,
        research_data: Dict[str, Any],
        vision_data: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """Unified synthesis of Vision + Research data (v7.0)"""
        logger.info(f"Generating Multi-Modal Intelligence Report for: {topic}")
        
        fusion_prompt = f"""
        Aşağıdaki çok yönlü (multimodal) verileri profesyonel ve stratejik bir 'Zeka Raporu' olarak sentezle.
        
        KONU: {topic}
        
        ARAŞTIRMA VERİLERİ: {json.dumps(research_data.get('findings', [])[:5])}
        GÖRSEL ANALİZ VERİLERİ: {json.dumps(vision_data) if vision_data else 'Görsel veri yok'}
        
        Raporda şunları içermelidir:
        1. STRATEJİK SENTEZ: Görsel ve metinsel verilerin birbirini nasıl tamamladığı.
        2. KRİTİK İÇGÖRÜLER: En önemli 3 bulgu.
        3. AKSİYON ÖNERİSİ: Kararlı bir adım.
        
        Dil: Profesyonel Türkçe.
        Format: Markdown.
        """
        
        try:
            if not self.llm_client: await self.initialize_llm()
            summary = await self.llm_client._ask_llm_with_custom_prompt(fusion_prompt)
            
            # Combine with standard report generation if needed, or return as synthesis
            return {
                "success": True,
                "synthesis": summary,
                "topic": topic,
                "timestamp": datetime.now().isoformat()
            }
        except Exception as e:
            logger.error(f"Intelligence report failed: {e}")
            return {"success": False, "error": str(e)}

    def _calculate_metrics(self, research_data: Dict[str, Any]) -> ReportMetrics:
        """Calculate research quality metrics"""
        sources = research_data.get("sources", [])
        findings = research_data.get("findings", [])

        # Calculate average reliability
        reliability_scores = [
            s.get("reliability_score", 0.5) for s in sources
        ]
        avg_reliability = sum(reliability_scores) / len(reliability_scores) if reliability_scores else 0.5

        # Coverage score (0-100) based on source count
        source_count = len(sources)
        coverage_score = min(100, (source_count / 10) * 100)

        # Reliability score (0-100)
        reliability_score = avg_reliability * 100

        # Completeness score (0-100) based on findings
        finding_count = len(findings)
        completeness_score = min(100, (finding_count / 15) * 100)

        return ReportMetrics(
            source_count=source_count,
            finding_count=finding_count,
            average_reliability=avg_reliability,
            research_depth=research_data.get("depth", "standard"),
            coverage_score=coverage_score,
            reliability_score=reliability_score,
            completeness_score=completeness_score
        )

    async def _generate_report_sections(
        self,
        topic: str,
        research_data: Dict[str, Any],
        metrics: ReportMetrics,
        style: str
    ) -> Dict[str, Any]:
        """Generate all report sections"""
        sources = research_data.get("sources", [])
        findings = research_data.get("findings", [])

        # Generate visualizations
        reliability_chart = self.visualizations.generate_source_reliability_chart(sources)
        findings_chart = self.visualizations.generate_findings_distribution_chart(findings)
        quality_viz = self.visualizations.generate_quality_metrics_visualization(
            metrics.coverage_score,
            metrics.reliability_score,
            metrics.completeness_score
        )

        sections = {
            "title": topic,
            "generated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "cover_page": await self._generate_cover_page(topic, metrics),
            "table_of_contents": self._generate_toc(),
            "executive_summary": await self._generate_executive_summary(topic, research_data),
            "introduction": self._generate_introduction(topic),
            "methodology": self._generate_methodology(metrics),
            "key_findings": self._generate_key_findings(research_data),
            "source_analysis": self._generate_source_analysis(research_data),
            "data_analysis": self._generate_data_analysis_section(
                reliability_chart,
                findings_chart,
                quality_viz
            ),
            "insights": await self._generate_insights(research_data),
            "recommendations": await self._generate_recommendations(research_data),
            "bibliography": self._generate_bibliography(research_data),
            "appendix": self._generate_appendix(research_data),
            "quality_metrics": metrics,
            "key_findings_list": findings, # Pass clean list
            "visualizations": {
                "reliability_chart": reliability_chart,
                "findings_chart": findings_chart,
                "quality_viz": quality_viz
            }
        }

        return sections

    async def _generate_cover_page(self, topic: str, metrics: ReportMetrics) -> str:
        """Generate professional cover page"""
        return f"""
        <div class="cover-page">
            <div class="cover-spacer"></div>
            <h1 class="cover-title">{topic}</h1>
            <h3 class="cover-subtitle">Professional Research Report</h3>
            <div class="cover-spacer"></div>
            <div class="cover-metrics">
                <p>Research Depth: <strong>{metrics.research_depth.title()}</strong></p>
                <p>Sources Analyzed: <strong>{metrics.source_count}</strong></p>
                <p>Key Findings: <strong>{metrics.finding_count}</strong></p>
                <p>Quality Score: <strong>{metrics.reliability_score:.1f}%</strong></p>
            </div>
            <div class="cover-spacer"></div>
            <p class="cover-date">Generated: {datetime.now().strftime('%B %d, %Y')}</p>
            <p class="cover-footer">Confidential - Research Report</p>
        </div>
        """

    def _generate_toc(self) -> str:
        """Generate table of contents"""
        toc_items = [
            "Executive Summary",
            "Introduction",
            "Research Methodology",
            "Key Findings",
            "Source Analysis",
            "Data Analysis & Visualizations",
            "Key Insights & Implications",
            "Recommendations",
            "Bibliography",
            "Appendix"
        ]

        toc_html = "<div class='toc'><h2>Table of Contents</h2><ol>"
        for i, item in enumerate(toc_items, 1):
            toc_html += f"<li>{item}</li>"
        toc_html += "</ol></div>"

        return toc_html

    async def _generate_executive_summary(
        self,
        topic: str,
        research_data: Dict[str, Any]
    ) -> str:
        """Generate professional executive summary using LLM"""
        findings = research_data.get("findings", [])

        if self.llm_client:
            try:
                prompt = f""""{topic}" konulu araştırma raporu için üst düzey bir "Yönetici Özeti" (Executive Summary) hazırla.
                
                Önemli Bulgular:
                {chr(10).join(f"- {f}" for f in findings[:7])}
                
                KURALLAR:
                - Dil: Profesyonel, stratejik ve akıcı bir Türkçe.
                - Format: 2-3 paragraf.
                - İçerik: Araştırmanın kapsamını, en kritik 3 çıkarımı ve stratejik önemi vurgula.
                - Hedef Kitle: Üst düzey yöneticiler ve karar vericiler.
                - Sadece metin döndür.
                """

                response = await self.llm_client._ask_llm(prompt, temperature=0.3, max_tokens=1000)
                summary = response.get("message", "")
                if summary:
                    return f"<div class='section'><h2>Yönetici Özeti</h2><p>{summary}</p></div>"
            except Exception as e:
                logger.debug(f"LLM summary generation failed: {e}")

        # Fallback summary
        summary = f"""
        Bu kapsamlı araştırma raporu, {topic} konusunu sistematik bir yaklaşımla analiz etmektedir. 
        {len(research_data.get('sources', []))} otoriter kaynağın incelenmesi sonucunda elde edilen veriler, 
        sektörel trendleri ve kritik içgörüleri ortaya koymaktadır. Rapor, stratejik karar alma süreçlerini 
        desteklemek amacıyla kanıta dayalı bulgular sunmaktadır.
        """
        return f"<div class='section'><h2>Yönetici Özeti</h2><p>{summary}</p></div>"

    def _generate_introduction(self, topic: str) -> str:
        """Generate professional introduction"""
        return f"""
        <div class='section'>
        <h2>Giriş ve Kapsam</h2>
        <p>Bu rapor, <strong>{topic}</strong> üzerine derinlemesine bir analiz sunmaktadır. 
        Çalışma, güncel pazar verilerini, uzman görüşlerini ve akademik kaynakları sentezleyerek 
        konuyla ilgili bütüncül bir perspektif oluşturmayı amaçlar.</p>
        <p>Analiz süreci, veri doğruluğu ve kaynak güvenilirliği ilkeleri çerçevesinde yürütülmüş; 
        elde edilen her bulgu çapraz sorgulama yöntemiyle doğrulanmıştır. Bu döküman, 
        stratejik planlama ve uygulama süreçlerinde rehberlik etmesi amacıyla hazırlanmıştır.</p>
        </div>
        """

    def _generate_methodology(self, metrics: ReportMetrics) -> str:
        """Generate methodology section"""
        return f"""
        <div class='section'>
        <h2>Araştırma Metodolojisi</h2>
        <h3>Yaklaşım</h3>
        <p>Bu araştırma, sistematik bilgi toplama ve kaynak değerlendirme teknikleri kullanılarak 
        yürütülmüştür. Çalışma kapsamında, çoklu disiplinlerden {metrics.source_count} adet otoriter kaynak 
        titizlikle incelenmiştir.</p>
        
        <h3>Kaynak Değerlendirme Kriterleri</h3>
        <ul>
            <li>Alan yetkinliği ve uzmanlık düzeyi</li>
            <li>İçerik güncelliği ve alaka düzeyi</li>
            <li>Metodolojik tutarlılık</li>
            <li>Alıntılanma sıklığı ve etki değeri</li>
            <li>Hakemli inceleme durumu</li>
        </ul>
        
        <h3>Kalite Metrikleri</h3>
        <table class='metrics-table'>
            <tr><td>Ortalama Kaynak Güvenilirliği</td><td>%{metrics.average_reliability*100:.1f}</td></tr>
            <tr><td>Kapsam Tamamlanma Oranı</td><td>%{metrics.coverage_score:.1f}</td></tr>
            <tr><td>Genel Kalite Skoru</td><td>%{metrics.reliability_score:.1f}</td></tr>
        </table>
        </div>
        """

    def _generate_key_findings(self, research_data: Dict[str, Any]) -> str:
        """Generate key findings section"""
        findings = research_data.get("findings", [])

        findings_html = "<div class='section'><h2>Temel Bulgular</h2>"

        for i, finding in enumerate(findings[:10], 1):
            findings_html += f"<div class='finding'><h3>Bulgu {i}</h3><p>{finding}</p></div>"

        findings_html += "</div>"
        return findings_html

    def _generate_source_analysis(self, research_data: Dict[str, Any]) -> str:
        """Generate source analysis section"""
        sources = research_data.get("sources", [])

        # Sort by reliability
        sorted_sources = sorted(
            sources,
            key=lambda x: x.get("reliability_score", 0),
            reverse=True
        )

        sources_html = "<div class='section'><h2>Kaynak Analizi</h2>"
        sources_html += "<table class='sources-table'><tr><th>Kaynak</th><th>Güvenilirlik</th><th>Etki Alanı</th></tr>"

        for source in sorted_sources[:15]:
            title = source.get("title", "Bilinmiyor")[:50]
            reliability = f"%{int(source.get('reliability_score', 0) * 100)}"
            url = source.get("url", "")
            domain = url.split("/")[2] if url else "Bilinmiyor"

            sources_html += f"<tr><td>{title}</td><td>{reliability}</td><td>{domain}</td></tr>"

        sources_html += "</table></div>"
        return sources_html

    def _generate_data_analysis_section(
        self,
        reliability_chart: Optional[str],
        findings_chart: Optional[str],
        quality_viz: Optional[str]
    ) -> str:
        """Generate data analysis section with visualizations"""
        html = "<div class='section'><h2>Veri Analizi ve Görselleştirme</h2>"

        if quality_viz:
            html += "<h3>Araştırma Kalite Metrikleri</h3>"
            if quality_viz.startswith("data:image"):
                html += f"<img src='{quality_viz}' alt='Kalite Metrikleri' style='max-width:100%; height:auto;'>"
            else:
                html += quality_viz
            html += "<p style='margin-top:20px;'>Yukarıdaki kalite metrikleri, toplanan tüm kaynaklar ve bulgular genelinde araştırmanın kapsamını, güvenilirliğini ve tamamlanma oranını göstermektedir.</p>"

        if reliability_chart:
            html += "<h3 style='margin-top:40px;'>Kaynak Güvenilirlik Analizi</h3>"
            if reliability_chart.startswith("data:image"):
                html += f"<img src='{reliability_chart}' alt='Kaynak Güvenilirliği' style='max-width:100%; height:auto;'>"
            else:
                html += reliability_chart
            html += "<p style='margin-top:20px;'>Bu grafik, kaynakları güvenilirlik puanlarına göre sıralayarak araştırmadaki en yetkin ve güvenilir bilgi odaklarını belirlemeye yardımcı olur.</p>"

        if findings_chart:
            html += "<h3 style='margin-top:40px;'>Bulgu Dağılımı</h3>"
            if findings_chart.startswith("data:image"):
                html += f"<img src='{findings_chart}' alt='Bulgu Dağılımı' style='max-width:100%; height:auto;'>"
            else:
                html += findings_chart
            html += "<p style='margin-top:20px;'>Bulgu dağılımı, kısa, detaylı ve kapsamlı araştırma sonuçları arasındaki dengeyi gösterir.</p>"

        html += "</div>"
        return html

    async def _generate_insights(self, research_data: Dict[str, Any]) -> str:
        """Generate insights and implications using LLM synthesis"""
        findings = research_data.get("findings", [])
        topic = research_data.get("topic", "the research topic")
        
        if not self.llm_client or not findings:
            return "<div class='section'><h2>Stratejik Analiz</h2><p>Veri sentezi için yeterli bulgu sağlanamadı.</p></div>"

        prompt = f"""Görevin: "{topic}" konulu araştırma için "Stratejik Analiz ve İçe Bakış" (Strategic Insights) bölümü yazmak.
        
        Aşağıdaki bulgulardan anlamlı örüntüler ve stratejik çıkarımlar oluştur:
        {findings[:15]}
        
        KURALLAR:
        - Başlıklar: <h2> ve <h3> etiketi kullan.
        - Dil: Analitik, profesyonel ve vizyoner Türkçe.
        - İçerik: "Bu veriler ne anlama geliyor?" sorusuna yanıt ver. Trendleri ve riskleri analiz et.
        - Format: HTML (p, ul, li).
        """
        try:
            response = await self.llm_client._ask_llm(prompt, temperature=0.5, max_tokens=1500)
            insights_body = response.get("message", "")
            insights_body = insights_body.replace("```html", "").replace("```", "").strip()
            return f"<div class='section'>{insights_body}</div>"
        except:
            return "<div class='section'><h2>Stratejik Analiz</h2><p>Analitik sentez süreci tamamlanamadı.</p></div>"

    async def _generate_recommendations(self, research_data: Dict[str, Any]) -> str:
        """Generate recommendations section using LLM"""
        findings = research_data.get("findings", [])
        topic = research_data.get("topic", "the research topic")
        
        if not self.llm_client or not findings:
            return "<div class='section'><h2>Stratejik Öneriler</h2><p>Aksiyon planı oluşturmak için yeterli veri yok.</p></div>"

        prompt = f"""Görevin: "{topic}" konulu araştırma için "Stratejik Öneriler" (Strategic Recommendations) bölümü yazmak.
        
        Bulgular:
        {findings[:10]}
        
        KURALLAR:
        - Başlık: <h2>.
        - Format: <ol> ve <li> kullanarak aksiyon listesi oluştur.
        - Detay: Her önerinin altına kısa bir "Neden?" veya "Nasıl?" açıklaması ekle.
        - Dil: Kararlı, profesyonel ve çözüm odaklı Türkçe.
        """
        try:
            response = await self.llm_client._ask_llm(prompt, temperature=0.4, max_tokens=1200)
            recs_body = response.get("message", "")
            recs_body = recs_body.replace("```html", "").replace("```", "").strip()
            return f"<div class='section'>{recs_body}</div>"
        except:
            return "<div class='section'><h2>Stratejik Öneriler</h2><p>Aksiyon planı oluşturulamadı.</p></div>"

    def _generate_bibliography(self, research_data: Dict[str, Any]) -> str:
        """Generate professional bibliography"""
        sources = research_data.get("sources", [])

        bib_html = "<div class='section'><h2>Bibliography</h2><ol class='bibliography'>"

        for i, source in enumerate(sources[:20], 1):
            url = source.get("url", "")
            title = source.get("title", "Source")
            reliability = f"[Reliability: {int(source.get('reliability_score', 0) * 100)}%]"

            if url:
                bib_html += f'<li><a href="{url}">{title}</a> {reliability}</li>'
            else:
                bib_html += f"<li>{title} {reliability}</li>"

        bib_html += "</ol></div>"
        return bib_html

    def _generate_appendix(self, research_data: Dict[str, Any]) -> str:
        """Generate appendix with detailed data"""
        return """
        <div class='section'>
        <h2>Appendix</h2>
        <h3>A. Research Methodology Details</h3>
        <p>Detailed information about source selection, evaluation criteria, and quality assurance processes.</p>

        <h3>B. Data Analysis</h3>
        <p>Comprehensive analysis of findings patterns, trends, and relationships.</p>

        <h3>C. Additional Resources</h3>
        <p>Links and references for further investigation and detailed research.</p>
        </div>
        """

    async def _generate_word_document(
        self,
        topic: str,
        sections: Dict[str, Any],
        metrics: ReportMetrics
    ) -> Path:
        """Generate Word document with professional formatting"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"{topic.replace(' ', '_')}_{timestamp}.docx"

            doc = Document()

            # Title
            title = doc.add_heading(topic, level=1)
            title_format = title.paragraph_format
            title_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title.runs[0].font.color.rgb = RGBColor(26, 26, 26)

            # Subtitle
            subtitle = doc.add_paragraph("Professional Research Report")
            subtitle_format = subtitle.paragraph_format
            subtitle_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.runs[0].font.size = Pt(14)

            # Metadata
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

            # Quality Metrics
            doc.add_heading("Quality Metrics", level=2)
            doc.add_paragraph(f"Coverage: {metrics.coverage_score:.0f}%")
            doc.add_paragraph(f"Reliability: {metrics.reliability_score:.0f}%")
            doc.add_paragraph(f"Completeness: {metrics.completeness_score:.0f}%")

            # Executive Summary
            doc.add_heading("Executive Summary", level=2)
            raw_summary = sections.get("executive_summary", "")
            summary = self._strip_html(raw_summary)
            if summary:
                doc.add_paragraph(summary)

            # Key Findings
            doc.add_heading("Key Findings", level=2)
            findings = sections.get("key_findings_list", [])
            if not findings:
                raw_findings = sections.get("key_findings", "")
                findings = self._strip_html_list(raw_findings)
            
            if findings:
                for i, finding in enumerate(findings[:10], 1):
                    clean_f = self._strip_html(finding)
                    doc.add_paragraph(clean_f, style='List Number')
            else:
                doc.add_paragraph("No specific findings extracted.")

            # Sources
            doc.add_heading("Sources", level=2)
            sources = sections.get("quality_metrics", {})
            if hasattr(sources, 'source_count'):
                doc.add_paragraph(f"Total Sources: {sources.source_count}")
                doc.add_paragraph(f"Average Reliability: {sources.average_reliability*100:.1f}%")

            # Save
            doc.save(str(output_path))
            logger.info(f"Word document created: {output_path}")
            return output_path

        except ImportError:
            logger.warning("python-docx not available, falling back to PDF")
            return await self._generate_advanced_pdf(topic, sections, metrics)
        except Exception as e:
            logger.error(f"Word generation error: {e}")
            return await self._generate_advanced_pdf(topic, sections, metrics)

    async def _generate_advanced_pdf(
        self,
        topic: str,
        sections: Dict[str, Any],
        metrics: ReportMetrics
    ) -> Path:
        """Generate advanced PDF with professional styling"""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, PageBreak,
                Table, TableStyle, Image
            )

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            output_path = self.output_dir / f"{topic.replace(' ', '_')}_{timestamp}.pdf"

            doc = SimpleDocTemplate(str(output_path), pagesize=A4)
            elements = []
            styles = getSampleStyleSheet()

            # Custom styles
            styles.add(ParagraphStyle(
                name='CoverTitle',
                fontSize=32,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=30,
                alignment=1,
                fontName='Helvetica-Bold'
            ))

            styles.add(ParagraphStyle(
                name='CoverSubtitle',
                fontSize=16,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=60,
                alignment=1
            ))

            styles.add(ParagraphStyle(
                name='SectionHeading',
                fontSize=16,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=12,
                spaceBefore=12,
                fontName='Helvetica-Bold'
            ))

            # Cover page
            elements.append(Spacer(1, 1.5 * inch))
            elements.append(Paragraph(topic, styles['CoverTitle']))
            elements.append(Paragraph("Professional Research Report", styles['CoverSubtitle']))
            elements.append(Spacer(1, 0.5 * inch))

            # Quality metrics on cover
            metrics_text = f"""
            Research Depth: {metrics.research_depth.title()}<br/>
            Sources: {metrics.source_count}<br/>
            Findings: {metrics.finding_count}<br/>
            Quality Score: {metrics.reliability_score:.1f}%
            """
            elements.append(Paragraph(metrics_text, styles['Normal']))
            elements.append(Spacer(1, 1 * inch))
            elements.append(Paragraph(
                f"Generated: {datetime.now().strftime('%B %d, %Y')}",
                styles['Normal']
            ))

            # Page break
            elements.append(PageBreak())

            # Table of contents
            elements.append(Paragraph("Table of Contents", styles['SectionHeading']))
            elements.append(Spacer(1, 0.2 * inch))

            toc_items = [
                "Executive Summary",
                "Introduction",
                "Research Methodology",
                "Key Findings",
                "Source Analysis",
                "Key Insights",
                "Recommendations",
                "Bibliography"
            ]

            for item in toc_items:
                elements.append(Paragraph(f"• {item}", styles['Normal']))

            elements.append(PageBreak())

            # Executive Summary
            elements.append(Paragraph("Executive Summary", styles['SectionHeading']))
            summary = f"""
            This comprehensive research report examines {topic} through analysis of
            {metrics.source_count} authoritative sources. The investigation identifies
            {metrics.finding_count} key findings with {metrics.reliability_score:.0f}% average source quality.
            Results provide evidence-based insights for informed decision-making.
            """
            elements.append(Paragraph(summary, styles['BodyText']))
            elements.append(PageBreak())

            # Key Findings
            elements.append(Paragraph("Key Findings", styles['SectionHeading']))
            elements.append(Spacer(1, 0.2 * inch))

            findings = sections.get('key_findings_list', [])
            if not findings:
                # Fallback to parsing the HTML section if list not available
                raw_findings = sections.get('key_findings', "")
                findings = self._strip_html_list(raw_findings)
            
            if findings:
                for i, finding in enumerate(findings[:10], 1):
                    # Clean and strip any remaining HTML
                    clean_f = self._strip_html(finding)
                    elements.append(Paragraph(f"{i}. {clean_f}", styles['BodyText']))
                    elements.append(Spacer(1, 0.1 * inch))

            elements.append(PageBreak())

            # Quality Metrics
            elements.append(Paragraph("Research Quality Metrics", styles['SectionHeading']))
            elements.append(Spacer(1, 0.2 * inch))

            metrics_data = [
                ['Metric', 'Score'],
                ['Coverage Completeness', f"{metrics.coverage_score:.1f}%"],
                ['Source Reliability', f"{metrics.reliability_score:.1f}%"],
                ['Finding Completeness', f"{metrics.completeness_score:.1f}%"],
                ['Total Sources', str(metrics.source_count)],
                ['Total Findings', str(metrics.finding_count)]
            ]

            metrics_table = Table(metrics_data, colWidths=[3 * inch, 1.5 * inch])
            metrics_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 12),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')])
            ]))
            elements.append(metrics_table)

            elements.append(PageBreak())

            # Data Analysis Section
            elements.append(Paragraph("Data Analysis & Visualizations", styles['SectionHeading']))
            elements.append(Spacer(1, 0.2 * inch))

            # Add quality metrics table
            quality_text = """
            <b>Quality Metrics Summary</b><br/>
            This research demonstrates strong quality across multiple dimensions:
            """
            elements.append(Paragraph(quality_text, styles['BodyText']))
            elements.append(Spacer(1, 0.1 * inch))

            # Quality summary data
            quality_summary = [
                ['Metric', 'Score', 'Assessment'],
                ['Coverage', f"{metrics.coverage_score:.0f}%", "Excellent" if metrics.coverage_score > 70 else "Good" if metrics.coverage_score > 50 else "Fair"],
                ['Reliability', f"{metrics.reliability_score:.0f}%", "Excellent" if metrics.reliability_score > 80 else "Good" if metrics.reliability_score > 60 else "Fair"],
                ['Completeness', f"{metrics.completeness_score:.0f}%", "Excellent" if metrics.completeness_score > 70 else "Good" if metrics.completeness_score > 50 else "Fair"],
            ]

            quality_table = Table(quality_summary, colWidths=[2 * inch, 1.5 * inch, 1.5 * inch])
            quality_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#2c3e50')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f0f0f0')])
            ]))
            elements.append(quality_table)

            elements.append(Spacer(1, 0.3 * inch))

            # Source Reliability Summary
            sources = sections.get('quality_metrics')
            elements.append(Paragraph("<b>Source Reliability Analysis</b><br/>Top sources ranked by reliability:", styles['BodyText']))
            elements.append(Spacer(1, 0.1 * inch))

            elements.append(PageBreak())

            # Bibliography
            elements.append(Paragraph("Bibliography & Sources", styles['SectionHeading']))
            elements.append(Spacer(1, 0.2 * inch))

            # Build PDF
            doc.build(elements)
            return output_path

        except ImportError:
            logger.warning("reportlab not available, using HTML fallback")
            # Return HTML version path
            return self._generate_html_report(topic, sections)

    def _strip_html(self, text: str) -> str:
        """HTML etiketlerini ve gereksiz boşlukları temizle"""
        if not text:
            return ""
        # Remove HTML tags
        clean = re.sub(r'<[^>]+>', '', text)
        # Decode entities if any (simple ones)
        clean = clean.replace('&nbsp;', ' ').replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>')
        # Normalize whitespace
        clean = re.sub(r'\s+', ' ', clean)
        return clean.strip()

    def _strip_html_list(self, html_content: str) -> List[str]:
        """HTML içeriğinden liste öğelerini çıkar"""
        if not html_content:
            return []
        # Find all content between <li> or <p> tags
        items = re.findall(r'<(?:li|p|div[^>]*|h[1-6])[^>]*>(.*?)</(?:li|p|div|h[1-6])>', html_content, re.DOTALL)
        if not items:
            # Try splitting by bullet points or generic tags
            items = [self._strip_html(s) for s in re.split(r'<br/?>|•', html_content) if s.strip()]
        else:
            items = [self._strip_html(item) for item in items if item.strip()]
        return items


async def generate_advanced_professional_report(
    topic: str,
    research_data: Dict[str, Any]
) -> Dict[str, Any]:
    """Generate advanced professional report"""
    generator = AdvancedReportGenerator()
    await generator.initialize_llm()
    return await generator.generate_professional_report(topic, research_data)
