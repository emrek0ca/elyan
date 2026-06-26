from __future__ import annotations

from pathlib import Path

import pytest

from runtime import capability_registry, state_store


def _dangerous_state() -> dict[str, object]:
    return state_store._ensure_defaults(  # type: ignore[attr-defined]
        {
            "account": {"dangerousAreaEnabled": True},
            "permissions": {"allow_destructive_tools": True},
        }
    )


def test_canvas_write_creates_pdf_with_table(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    import actions.canvas_write as canvas_write

    monkeypatch.setattr(canvas_write, "_workspace_root", lambda: tmp_path)

    result = capability_registry.run_capability(
        "canvas_write",
        {
            "_confirmed": True,
            "outputPath": str(tmp_path / "canvas-report.pdf"),
            "outputFormat": "pdf",
            "title": "Canvas Raporu",
            "blocks": [
                {"kind": "text", "text": "Elyan Canvas"},
                {
                    "kind": "table",
                    "title": "Durum Tablosu",
                    "headers": ["Alan", "Değer"],
                    "rows": [["Model", "Elyan"], ["Durum", "Hazır"]],
                },
            ],
        },
        _dangerous_state(),
    )

    assert result["ok"] is True
    output_path = Path(result["result"]["outputPath"])
    assert output_path.exists()
    assert output_path.stat().st_size > 0
    assert result["result"]["blockCount"] == 2


def test_canvas_write_creates_png_from_source_image(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    import actions.canvas_write as canvas_write
    from PIL import Image

    monkeypatch.setattr(canvas_write, "_workspace_root", lambda: tmp_path)

    source_path = tmp_path / "source-image.png"
    Image.new("RGB", (320, 180), color=(220, 210, 195)).save(source_path)

    result = capability_registry.run_capability(
        "canvas_write",
        {
            "_confirmed": True,
            "sourcePath": str(source_path),
            "outputPath": str(tmp_path / "canvas-preview.png"),
            "outputFormat": "png",
            "title": "Canvas Önizleme",
        },
        _dangerous_state(),
    )

    assert result["ok"] is True
    output_path = Path(result["result"]["outputPath"])
    assert output_path.exists()
    assert output_path.suffix == ".png"
    assert output_path.stat().st_size > 0


def test_canvas_write_creates_png_with_chart_block(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    import actions.canvas_write as canvas_write
    from PIL import Image

    monkeypatch.setattr(canvas_write, "_workspace_root", lambda: tmp_path)

    result = capability_registry.run_capability(
        "canvas_write",
        {
            "_confirmed": True,
            "outputPath": str(tmp_path / "canvas-chart.png"),
            "outputFormat": "png",
            "title": "Grafik Kanvası",
            "blocks": [
                {"kind": "text", "text": "Satış Trendi"},
                {
                    "kind": "chart",
                    "chartType": "line",
                    "title": "Aylık Satış",
                    "labels": ["Ocak", "Şubat", "Mart", "Nisan"],
                    "values": [12, 19, 15, 24],
                    "width": 960,
                    "height": 540,
                },
            ],
        },
        _dangerous_state(),
    )

    assert result["ok"] is True
    output_path = Path(result["result"]["outputPath"])
    assert output_path.exists()
    assert output_path.stat().st_size > 0
    with Image.open(output_path) as image:
        assert image.size[0] >= 800


def test_document_write_accepts_table_blocks(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    import actions.document_write as document_write
    from docx import Document

    monkeypatch.setattr(document_write, "_workspace_root", lambda: tmp_path)

    result = capability_registry.run_capability(
        "document_write",
        {
            "_confirmed": True,
            "outputPath": str(tmp_path / "table-report.docx"),
            "title": "Tablo Raporu",
            "blocks": [
                {"kind": "text", "text": "Başlık"},
                {
                    "kind": "table",
                    "title": "Özet Tablosu",
                    "headers": ["Alan", "Değer"],
                    "rows": [["A", "1"], ["B", "2"]],
                },
            ],
        },
        _dangerous_state(),
    )

    assert result["ok"] is True
    output_path = Path(result["result"]["outputPath"])
    assert output_path.exists()
    document = Document(str(output_path))
    assert len(document.tables) == 1


def test_document_write_accepts_chart_blocks(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    import actions.document_write as document_write
    from docx import Document

    monkeypatch.setattr(document_write, "_workspace_root", lambda: tmp_path)

    result = capability_registry.run_capability(
        "document_write",
        {
            "_confirmed": True,
            "outputPath": str(tmp_path / "chart-report.docx"),
            "title": "Grafik Raporu",
            "blocks": [
                {"kind": "text", "text": "Aylık trend"},
                {
                    "kind": "chart",
                    "chartType": "bar",
                    "title": "Kullanım",
                    "labels": ["A", "B", "C"],
                    "values": [2, 5, 3],
                },
            ],
        },
        _dangerous_state(),
    )

    assert result["ok"] is True
    output_path = Path(result["result"]["outputPath"])
    assert output_path.exists()
    document = Document(str(output_path))
    assert len(document.inline_shapes) >= 1


def test_presentation_write_accepts_chart_blocks(monkeypatch: pytest.MonkeyPatch, tmp_path: Path) -> None:
    import actions.presentation_write as presentation_write
    from pptx import Presentation

    monkeypatch.setattr(presentation_write, "_workspace_root", lambda: tmp_path)

    result = capability_registry.run_capability(
        "presentation_write",
        {
            "_confirmed": True,
            "outputPath": str(tmp_path / "chart-deck.pptx"),
            "title": "Grafik Sunumu",
            "blocks": [
                {
                    "kind": "chart",
                    "chartType": "scatter",
                    "title": "Dağılım",
                    "labels": ["1", "2", "3", "4"],
                    "values": [2.5, 3.0, 1.8, 4.4],
                }
            ],
        },
        _dangerous_state(),
    )

    assert result["ok"] is True
    output_path = Path(result["result"]["outputPath"])
    assert output_path.exists()
    presentation = Presentation(str(output_path))
    assert len(presentation.slides) >= 2
