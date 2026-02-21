"""Word Document Tools - Read and Write .docx files"""

import asyncio
from pathlib import Path
from typing import Any
from utils.logger import get_logger
from security.validator import validate_path
from config.settings import HOME_DIR, DESKTOP

logger = get_logger("office.word")

# Maximum file size (10MB)
MAX_FILE_SIZE = 10 * 1024 * 1024


async def read_word(path: str, max_chars: int = 10000) -> dict[str, Any]:
    """Read content from a Word document (.docx)

    Args:
        path: Path to the Word document
        max_chars: Maximum characters to return (default 10000)
    """
    try:
        # Validate path
        is_valid, error, _ = validate_path(path)
        if not is_valid:
            return {"success": False, "error": error}

        file_path = Path(path).expanduser().resolve()

        if not file_path.exists():
            return {"success": False, "error": f"Dosya bulunamadı: {file_path.name}"}

        if not file_path.suffix.lower() in [".docx", ".doc"]:
            return {"success": False, "error": "Sadece .docx dosyaları destekleniyor"}

        # Check file size
        if file_path.stat().st_size > MAX_FILE_SIZE:
            return {"success": False, "error": "Dosya çok büyük (max 10MB)"}

        # Import here to avoid startup delay if not used
        try:
            from docx import Document
        except ImportError:
            return {"success": False, "error": "python-docx kurulu değil. 'pip install python-docx' çalıştırın."}

        # Run in thread pool to avoid blocking
        def _read():
            doc = Document(str(file_path))
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)

            # Also get text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        tables_text.append(row_text)

            return paragraphs, tables_text, len(doc.paragraphs), len(doc.tables)

        loop = asyncio.get_event_loop()
        paragraphs, tables_text, para_count, table_count = await loop.run_in_executor(None, _read)

        # Combine content
        content = "\n\n".join(paragraphs)
        if tables_text:
            content += "\n\n--- Tablolar ---\n" + "\n".join(tables_text)

        # Truncate if needed
        truncated = False
        if len(content) > max_chars:
            content = content[:max_chars]
            truncated = True

        logger.info(f"Read Word file: {file_path.name}, {len(content)} chars")

        return {
            "success": True,
            "path": str(file_path),
            "filename": file_path.name,
            "content": content,
            "paragraph_count": para_count,
            "table_count": table_count,
            "truncated": truncated
        }

    except Exception as e:
        logger.error(f"Read Word error: {e}")
        return {"success": False, "error": str(e)}


async def write_word(
    path: str = None,
    content: str = "",
    title: str = None,
    paragraphs: list = None
) -> dict[str, Any]:
    """Create or write to a Word document (.docx)

    Args:
        path: Path for the document (default: Desktop/document.docx)
        content: Text content to write (will be split by newlines into paragraphs)
        title: Optional document title (added as heading)
        paragraphs: Optional list of paragraphs (alternative to content string)
    """
    try:
        # Default path
        if not path:
            path = str(DESKTOP / "belge.docx")

        # Ensure .docx extension
        file_path = Path(path).expanduser().resolve()
        if not file_path.suffix.lower() == ".docx":
            file_path = file_path.with_suffix(".docx")

        # Validate path
        is_valid, error, _ = validate_path(str(file_path))
        if not is_valid:
            return {"success": False, "error": error}

        content_text = str(content or "").strip()
        paragraph_items: list[str] = []
        if isinstance(paragraphs, list):
            for para in paragraphs:
                para_text = str(para or "").strip()
                if para_text:
                    paragraph_items.append(para_text)

        if not str(title or "").strip() and not content_text and not paragraph_items:
            return {"success": False, "error": "Word dosyası için yazılacak içerik boş."}

        try:
            from docx import Document
            from docx.shared import Inches, Pt
        except ImportError:
            return {"success": False, "error": "python-docx kurulu değil. 'pip install python-docx' çalıştırın."}

        def _write():
            doc = Document()

            # Add title if provided
            if title:
                doc.add_heading(title, 0)

            # Add content
            if paragraph_items:
                for para in paragraph_items:
                    doc.add_paragraph(para)
            elif content_text:
                # Split by double newlines for paragraphs
                for para in content_text.split("\n\n"):
                    if para.strip():
                        doc.add_paragraph(para.strip())

            # Create parent directory if needed
            file_path.parent.mkdir(parents=True, exist_ok=True)

            doc.save(str(file_path))
            return True

        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, _write)

        logger.info(f"Created Word file: {file_path.name}")

        return {
            "success": True,
            "path": str(file_path),
            "filename": file_path.name
        }

    except Exception as e:
        logger.error(f"Write Word error: {e}")
        return {"success": False, "error": str(e)}
