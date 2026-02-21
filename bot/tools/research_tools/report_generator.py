"""
Professional Research Report Generator
Generates formatted reports with text, tables, and charts
"""

import asyncio
from typing import Dict, List, Optional, Any
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass
from utils.logger import get_logger

logger = get_logger("report_generator")


@dataclass
class ReportSection:
    """Report section"""
    title: str
    content: str
    subsections: List["ReportSection"] = None

    def __post_init__(self):
        if self.subsections is None:
            self.subsections = []


class ProfessionalReportGenerator:
    """Generates professional research reports in multiple formats"""

    def __init__(self):
        self.output_dir = Path.home() / "Desktop" / "Research Reports"
        self.output_dir.mkdir(parents=True, exist_ok=True)

    async def generate_from_research(
        self,
        topic: str,
        research_data: Dict[str, Any],
        format: str = "pdf"
    ) -> Dict[str, Any]:
        """Generate report from research data"""
        try:
            report_name = self._sanitize_filename(topic)
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"{report_name}_{timestamp}"

            # Prepare report structure
            report = await self._build_report_structure(topic, research_data)

            # Generate output
            if format.lower() == "pdf":
                output_path = await self._generate_pdf(filename, report)
            elif format.lower() == "docx":
                output_path = await self._generate_docx(filename, report)
            else:
                output_path = await self._generate_html(filename, report)

            logger.info(f"Report generated: {output_path}")

            return {
                "success": True,
                "filename": output_path.name,
                "path": str(output_path),
                "format": format,
                "size_mb": output_path.stat().st_size / (1024 * 1024)
            }

        except Exception as e:
            logger.error(f"Report generation failed: {e}")
            return {
                "success": False,
                "error": str(e)
            }

    async def _build_report_structure(
        self,
        topic: str,
        research_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Build complete report structure"""
        sources = research_data.get("sources", [])
        findings = research_data.get("findings", [])
        summary = research_data.get("summary", "")

        return {
            "metadata": {
                "title": f"Research Report: {topic}",
                "topic": topic,
                "generated": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                "author": "Research Assistant",
            },
            "executive_summary": summary or self._generate_summary(findings),
            "introduction": self._generate_introduction(topic),
            "sections": self._generate_sections(findings),
            "sources": self._format_sources(sources),
            "statistics": self._generate_statistics(research_data),
            "tables": self._generate_tables(research_data),
        }

    async def _generate_pdf(self, filename: str, report: Dict[str, Any]) -> Path:
        """Generate PDF report"""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
            from reportlab.pdfgen import canvas

            output_path = self.output_dir / f"{filename}.pdf"

            # Create PDF document
            doc = SimpleDocTemplate(str(output_path), pagesize=A4)
            elements = []
            styles = getSampleStyleSheet()

            # Custom styles
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                textColor=colors.HexColor('#1a1a1a'),
                spaceAfter=30,
                alignment=1
            )

            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#2c3e50'),
                spaceAfter=12,
                spaceBefore=12
            )

            body_style = ParagraphStyle(
                'CustomBody',
                parent=styles['BodyText'],
                fontSize=11,
                alignment=4,
                spaceAfter=12
            )

            # Title page
            metadata = report.get("metadata", {})
            elements.append(Spacer(1, 1.5 * inch))
            elements.append(Paragraph(metadata.get("title", "Research Report"), title_style))
            elements.append(Spacer(1, 0.3 * inch))
            elements.append(Paragraph(
                f"Generated: {metadata.get('generated')}",
                styles['Normal']
            ))
            elements.append(PageBreak())

            # Executive Summary
            elements.append(Paragraph("Executive Summary", heading_style))
            summary = report.get("executive_summary", "")
            elements.append(Paragraph(summary, body_style))
            elements.append(Spacer(1, 0.3 * inch))

            # Introduction
            elements.append(PageBreak())
            elements.append(Paragraph("Introduction", heading_style))
            intro = report.get("introduction", "")
            elements.append(Paragraph(intro, body_style))
            elements.append(Spacer(1, 0.3 * inch))

            # Main sections
            for section in report.get("sections", []):
                elements.append(Paragraph(section.get("title", ""), heading_style))
                elements.append(Paragraph(section.get("content", ""), body_style))
                elements.append(Spacer(1, 0.2 * inch))

            # Statistics
            if report.get("statistics"):
                elements.append(PageBreak())
                elements.append(Paragraph("Key Statistics", heading_style))
                stats_data = report["statistics"]
                elements.append(Paragraph(stats_data, body_style))
                elements.append(Spacer(1, 0.2 * inch))

            # Sources
            elements.append(PageBreak())
            elements.append(Paragraph("Sources and References", heading_style))
            sources = report.get("sources", [])
            if isinstance(sources, list):
                for i, source in enumerate(sources, 1):
                    elements.append(Paragraph(
                        f"{i}. {source}",
                        body_style
                    ))

            # Build PDF
            doc.build(elements)

            return output_path

        except ImportError:
            logger.warning("reportlab not installed, using fallback HTML")
            return await self._generate_html(filename, report)

    async def _generate_docx(self, filename: str, report: Dict[str, Any]) -> Path:
        """Generate DOCX report"""
        try:
            from docx import Document
            from docx.shared import Pt, RGBColor, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            output_path = self.output_dir / f"{filename}.docx"
            doc = Document()

            # Title
            metadata = report.get("metadata", {})
            title = doc.add_heading(metadata.get("title", "Research Report"), level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Metadata
            meta_para = doc.add_paragraph()
            meta_para.add_run(f"Generated: {metadata.get('generated')}\n")
            meta_para.add_run(f"Topic: {metadata.get('topic')}")
            meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()  # Spacing

            # Executive Summary
            doc.add_heading("Executive Summary", level=2)
            summary = report.get("executive_summary", "")
            doc.add_paragraph(summary)

            # Introduction
            doc.add_heading("Introduction", level=2)
            intro = report.get("introduction", "")
            doc.add_paragraph(intro)

            # Main sections
            for section in report.get("sections", []):
                doc.add_heading(section.get("title", ""), level=2)
                doc.add_paragraph(section.get("content", ""))

            # Statistics
            if report.get("statistics"):
                doc.add_heading("Key Statistics", level=2)
                doc.add_paragraph(report["statistics"])

            # Sources
            doc.add_heading("Sources and References", level=2)
            sources = report.get("sources", [])
            if isinstance(sources, list):
                for i, source in enumerate(sources, 1):
                    doc.add_paragraph(source, style='List Number')

            doc.save(str(output_path))
            return output_path

        except ImportError:
            logger.warning("python-docx not installed, using fallback HTML")
            return await self._generate_html(filename, report)

    async def _generate_html(self, filename: str, report: Dict[str, Any]) -> Path:
        """Generate HTML report"""
        output_path = self.output_dir / f"{filename}.html"

        metadata = report.get("metadata", {})
        html_content = f"""
<!DOCTYPE html>
<html lang="tr">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{metadata.get('title', 'Research Report')}</title>
    <style>
        * {{
            margin: 0;
            padding: 0;
            box-sizing: border-box;
        }}

        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            color: #1a1a1a;
            background: #f8f8f8;
        }}

        .container {{
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            background: white;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }}

        .title-page {{
            text-align: center;
            border-bottom: 2px solid #2c3e50;
            padding-bottom: 30px;
            margin-bottom: 40px;
        }}

        h1 {{
            font-size: 2.5em;
            color: #1a1a1a;
            margin-bottom: 20px;
            font-weight: 600;
        }}

        h2 {{
            font-size: 1.8em;
            color: #2c3e50;
            margin-top: 40px;
            margin-bottom: 20px;
            padding-bottom: 10px;
            border-bottom: 1px solid #ecf0f1;
        }}

        h3 {{
            font-size: 1.2em;
            color: #34495e;
            margin-top: 20px;
            margin-bottom: 10px;
        }}

        p {{
            margin-bottom: 15px;
            text-align: justify;
        }}

        .metadata {{
            color: #7f8c8d;
            font-size: 0.95em;
            margin: 10px 0;
        }}

        .executive-summary {{
            background: #ecf0f1;
            padding: 20px;
            border-left: 4px solid #3498db;
            margin: 20px 0;
            border-radius: 4px;
        }}

        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 20px 0;
        }}

        th {{
            background: #34495e;
            color: white;
            padding: 12px;
            text-align: left;
            font-weight: 600;
        }}

        td {{
            padding: 12px;
            border-bottom: 1px solid #ecf0f1;
        }}

        tr:hover {{
            background: #f8f9fa;
        }}

        .sources {{
            margin-top: 40px;
            border-top: 2px solid #ecf0f1;
            padding-top: 20px;
        }}

        .source-item {{
            margin-bottom: 15px;
            padding-left: 20px;
        }}

        .source-item::before {{
            content: '▪';
            margin-left: -15px;
            margin-right: 10px;
            color: #3498db;
        }}

        .source-item a {{
            color: #3498db;
            text-decoration: none;
        }}

        .source-item a:hover {{
            text-decoration: underline;
        }}

        .statistics {{
            background: #f8f9fa;
            padding: 20px;
            border-radius: 4px;
            margin: 20px 0;
        }}

        .footer {{
            text-align: center;
            color: #7f8c8d;
            font-size: 0.9em;
            margin-top: 40px;
            padding-top: 20px;
            border-top: 1px solid #ecf0f1;
        }}
    </style>
</head>
<body>
    <div class="container">
        <div class="title-page">
            <h1>{metadata.get('title', 'Research Report')}</h1>
            <div class="metadata">
                <p>Generated: {metadata.get('generated')}</p>
                <p>Topic: {metadata.get('topic')}</p>
            </div>
        </div>

        <div class="executive-summary">
            <h2 style="margin-top: 0; border: none;">Executive Summary</h2>
            <p>{report.get('executive_summary', '')}</p>
        </div>

        <h2>Introduction</h2>
        <p>{report.get('introduction', '')}</p>

        <h2>Key Findings</h2>
        {''.join(f"<h3>{s.get('title', '')}</h3><p>{s.get('content', '')}</p>" for s in report.get('sections', []))}

        <div class="statistics">
            <h2 style="margin-top: 0;">Statistics</h2>
            <p>{report.get('statistics', '')}</p>
        </div>

        <div class="sources">
            <h2>Sources and References</h2>
            {''.join(f'<div class="source-item">{source}</div>' for source in report.get('sources', []))}
        </div>

        <div class="footer">
            <p>This report was automatically generated by Research Assistant</p>
        </div>
    </div>
</body>
</html>
"""

        output_path.write_text(html_content, encoding='utf-8')
        return output_path

    def _sanitize_filename(self, filename: str) -> str:
        """Sanitize filename for filesystem"""
        import re
        filename = re.sub(r'[<>:"/\\|?*]', '', filename)
        filename = filename.replace(' ', '_')
        return filename[:50]

    def _generate_summary(self, findings: List[str]) -> str:
        """Generate executive summary from findings"""
        if not findings:
            return "Research findings are presented in detail in the following sections."

        summary = "This research report presents the following key findings:\n\n"
        for i, finding in enumerate(findings[:5], 1):
            summary += f"{i}. {finding}\n"
        return summary

    def _generate_introduction(self, topic: str) -> str:
        """Generate introduction section"""
        return (
            f"This report presents a comprehensive analysis of '{topic}'. "
            "The research was conducted through systematic information gathering, "
            "source evaluation, and synthesis of relevant findings. "
            "The following sections detail the key aspects, sources, and conclusions."
        )

    def _generate_sections(self, findings: List[str]) -> List[Dict[str, str]]:
        """Generate main report sections from findings"""
        sections = []
        categories = {
            "Overview": findings[:2] if findings else [],
            "Analysis": findings[2:4] if len(findings) > 2 else [],
            "Key Insights": findings[4:6] if len(findings) > 4 else [],
            "Implications": findings[6:] if len(findings) > 6 else [],
        }

        for title, items in categories.items():
            if items:
                content = "\n".join(f"• {item}" for item in items)
                sections.append({
                    "title": title,
                    "content": content
                })

        return sections

    def _format_sources(self, sources: List[Dict[str, Any]]) -> List[str]:
        """Format sources for report"""
        formatted = []
        for source in sources:
            url = source.get("url", "")
            title = source.get("title", "Source")
            reliability = source.get("reliability_score", 0)

            reliability_text = ""
            if reliability > 0.8:
                reliability_text = " (Highly Reliable)"
            elif reliability > 0.6:
                reliability_text = " (Reliable)"
            elif reliability > 0.4:
                reliability_text = " (Moderate Reliability)"

            if url:
                formatted.append(f'<a href="{url}">{title}</a>{reliability_text}')
            else:
                formatted.append(f'{title}{reliability_text}')

        return formatted

    def _generate_statistics(self, research_data: Dict[str, Any]) -> str:
        """Generate statistics section"""
        sources_count = len(research_data.get("sources", []))
        findings_count = len(research_data.get("findings", []))
        depth = research_data.get("depth", "unknown")

        return (
            f"Sources Analyzed: {sources_count}\n"
            f"Key Findings: {findings_count}\n"
            f"Research Depth: {depth}\n"
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        )

    def _generate_tables(self, research_data: Dict[str, Any]) -> List[List[str]]:
        """Generate tables for report"""
        tables = []

        # Sources table
        sources = research_data.get("sources", [])
        if sources:
            source_table = [["Source", "Reliability", "Status"]]
            for source in sources[:10]:
                reliability = f"{int(source.get('reliability_score', 0) * 100)}%"
                status = "Fetched" if source.get("fetched") else "Pending"
                source_table.append([source.get("title", "")[:40], reliability, status])
            tables.append(source_table)

        return tables


# Global instance
_report_generator: Optional[ProfessionalReportGenerator] = None


def get_report_generator() -> ProfessionalReportGenerator:
    """Get or create report generator"""
    global _report_generator
    if _report_generator is None:
        _report_generator = ProfessionalReportGenerator()
    return _report_generator


async def create_research_report(
    topic: str,
    research_data: Dict[str, Any],
    format: str = "pdf"
) -> Dict[str, Any]:
    """Create research report from research data"""
    generator = get_report_generator()
    return await generator.generate_from_research(topic, research_data, format)
